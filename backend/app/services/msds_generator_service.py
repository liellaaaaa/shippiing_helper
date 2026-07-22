"""
MSDS 自动化生成服务。

从 customs_codes.json 获取产品数据，从旧 MSDS 文件生成新 MSDS。
"""
import json
import re
import os
from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO

from app.core.config import CUSTOMS_CODES_JSON, MSDS_DIR, TEMPLATES


# 项目根目录
ROOT = Path(__file__).parent.parent.parent.parent
INGREDIENT_MAPPING_JSON = str(ROOT / "references" / "ingredient_mapping.json")
APPEARANCE_COLOR_MAPPING_JSON = str(ROOT / "references" / "appearance_color_mapping.json")
PRODUCTS_NAME_MAPPING_JSON = str(ROOT / "references" / "products_name_mapping.json")
ION_TYPE_MAPPING_JSON = str(ROOT / "references" / "ion_type_mapping.json")
MSDS_ENGLISH_TEMPLATE_JSON = str(ROOT / "references" / "msds_english_template.json")

# 英文 MSDS 参考文件目录
ENGLISH_MSDS_DIR = str(ROOT / "references" / "MSDS")

# 英文 MSDS 模板文件路径
ENGLISH_MSDS_TEMPLATE = str(ROOT / "references" / "MSDS" / "模板 英文MSDS.docx")

# 英文 Section 标题（16 节）
SECTION_TITLES_EN = [
    "MATERIAL SAFETY DATA SHEET",
    "1. Identification",
    "2. Hazards Identification",
    "3. Main Component",
    "4. First Aid Measures",
    "5. Fire Protection Measures",
    "6. Accidental Release Measures",
    "7. Handling and Storage",
    "8. Exposure Controls/Personal Protection",
    "9. Physical and Chemical Properties",
    "10. Stability and Reactivity",
    "11. Toxicological Information",
    "12. Ecological Information",
    "13. Disposal Considerations",
    "14. Transport Information",
    "15. Regulatory Information",
    "16. Additional Information",
]

# 英文理化特性 Keys
PHYSLICOCHEMICAL_KEYS_EN = {
    "appearance": ["Appearance:", "Appearance/Physical form:"],
    "ionicity": ["Ionicity:", "Ionic nature:"],
    "ph": ["pH:", "pH value:"],
    "melting_point": ["Melting point:", "Melting Point:"],
    "boiling_point": ["Boiling point:", "Boiling Point:", "Boiling point(℃):"],
    "density": ["Density(℃):", "Density:", "Relative density:"],
    "flash_point": ["Flash Point:", "Flame point:"],
    "solubility": ["Solubility:", "Solubility in water:"],
    "lower_explosive": ["Lower explosive limit:", "Lower explosive limit"],
    "upper_explosive": ["Upper explosive limit:", "Upper explosive limit"],
    "steam_pressure": ["Steam pressure:", "Vapor pressure:"],
}

# 英文成分表列头
COMPOSITION_HEADERS_EN = ["Component", "CAS NO.", "W/W %"]


class MSDSGeneratorService:
    def __init__(self):
        self._customs_codes: list[dict] = []
        self._ingredient_map: list[dict] = []
        self._appearance_map: dict[str, str] = {}
        self._products_name_map: dict[str, str] = {}
        self._ion_type_map: dict[str, str] = {}
        # 英文 MSDS 翻译模板（从 msds_english_template.json 加载）
        self._msds_template: dict = {}
        # 翻译映射（按 key 长度降序排列，用于最长前缀匹配）
        self._translation_mapping: list[tuple[str, str]] = []
        # 英文 MSDS 结构缓存
        self._english_structure_cache: dict = {}
        self._load_data()
        self._load_msds_template()

    def _load_data(self):
        """加载所有数据文件"""
        # 加载 customs_codes.json
        if Path(CUSTOMS_CODES_JSON).exists():
            with open(CUSTOMS_CODES_JSON, "r", encoding="utf-8") as f:
                self._customs_codes = json.load(f)

        # 加载 ingredient_mapping.json
        if Path(INGREDIENT_MAPPING_JSON).exists():
            with open(INGREDIENT_MAPPING_JSON, "r", encoding="utf-8") as f:
                data = json.load(f)
                self._ingredient_map = data.get("mappings", [])

        # 加载 appearance_color_mapping.json
        if Path(APPEARANCE_COLOR_MAPPING_JSON).exists():
            with open(APPEARANCE_COLOR_MAPPING_JSON, "r", encoding="utf-8") as f:
                data = json.load(f)
                for item in data.get("mappings", []):
                    self._appearance_map[item["cn"]] = item["en"]

        # 加载 products_name_mapping.json
        if Path(PRODUCTS_NAME_MAPPING_JSON).exists():
            with open(PRODUCTS_NAME_MAPPING_JSON, "r", encoding="utf-8") as f:
                data = json.load(f)
                for item in data.get("mappings", []):
                    self._products_name_map[item["cn"]] = item["en"]

        # 加载 ion_type_mapping.json
        if Path(ION_TYPE_MAPPING_JSON).exists():
            with open(ION_TYPE_MAPPING_JSON, "r", encoding="utf-8") as f:
                data = json.load(f)
                for item in data.get("mappings", []):
                    self._ion_type_map[item["cn"]] = item["en"]

    def _load_msds_template(self):
        """加载 msds_english_template.json 并构建翻译映射"""
        if not Path(MSDS_ENGLISH_TEMPLATE_JSON).exists():
            return
        with open(MSDS_ENGLISH_TEMPLATE_JSON, "r", encoding="utf-8") as f:
            self._msds_template = json.load(f)
        self._build_translation_mapping()

    def _normalize(self, text: str) -> str:
        """
        归一化文本：全角→半角、去除多余空格、标点标准化。
        用于模糊匹配阶段，解决 encoding/空格/标点差异导致的匹配失败。
        """
        import unicodedata
        import re

        # 1. Unicode 归一化（NFKC 规范化，将全角字符转为兼容半角）
        text = unicodedata.normalize('NFKC', text)

        # 2. 全角数字 → 半角（UNICODE 范围）
        text = text.translate(str.maketrans(
            '０１２３４５６７８９',
            '0123456789'
        ))

        # 3. 中文顿号 → 英文逗号+空格（解决 "土壤、河流、下水道" 无法匹配 "土壤, 河流, 下水道" 的问题）
        text = text.replace('、', ', ')

        # 4. 去除多余连续空格（保留单词间单空格）
        text = re.sub(r'\s+', ' ', text)

        # 5. 去除前后空格
        text = text.strip()

        return text

    def _ensure_argos_model(self):
        """确保 argostranslate 中文→英文模型已下载"""
        try:
            from argostranslate import translate as t
            # 检查是否已有可用翻译
            tr = t.get_translation_from_codes('zh', 'en')
            if tr is not None:
                return
            # 尝试下载
            t.package.update_package_index()
            available = t.package.get_available_packages()
            for pkg in available:
                if getattr(pkg, 'code', None) == 'translate-zh_en':
                    pkg.install()
                    break
        except Exception:
            pass

    def _translate_with_argos(self, text_cn: str) -> str:
        """使用 translatepy 翻译中文→英文（第三层 fallback）"""
        try:
            from translatepy import Translator
            t = Translator()
            result = t.translate(text_cn, "English", "Chinese")
            if result and hasattr(result, 'result'):
                return result.result
            elif isinstance(result, str):
                return result
            return text_cn
        except Exception:
            return text_cn

    def _build_translation_mapping(self):
        """从 msds_english_template.json 提取所有翻译条目，按 key 长度降序排列"""
        if not self._msds_template:
            return
        template = self._msds_template.get("template", {})

        all_keys: list[tuple[str, str]] = []

        # 1. 章节标题
        section_titles = template.get("section_titles", {})
        for cn, en in section_titles.items():
            if cn != "MATERIAL SAFETY DATA SHEET":  # 主标题单独处理
                all_keys.append((cn, en))

        # 2. 各节字段（section1_fields ... section16_fields）
        for key in template:
            if key.startswith("section") and key.endswith("_fields"):
                for cn, en in template[key].items():
                    all_keys.append((cn, en))

        # 3. 固定短语 fixed_texts（同时注册原始 key 和 normalize 后的 key）
        fixed_texts = template.get("fixed_texts", {})
        for cn, en in fixed_texts.items():
            all_keys.append((cn, en))
            # 同时注册 normalize 后的 key（用于模糊匹配）
            cn_norm = self._normalize(cn)
            if cn_norm != cn:
                all_keys.append((cn_norm, en))

        # 按 key 长度降序排列（避免短 key 先匹配导致长 key 无法匹配）
        all_keys.sort(key=lambda x: len(x[0]), reverse=True)

        # 添加全角标点符号转换（放在最后，所有字段翻译之后）
        punctuation_map = [
            ("：", ": "),   # 全角冒号 → 英文冒号+空格
            ("，", ", "),   # 全角逗号 → 英文逗号+空格
            ("。", ". "),   # 全角句号 → 英文句号+空格
        ]
        for cn, en in punctuation_map:
            all_keys.append((cn, en))

        self._translation_mapping = all_keys

    def _translate_paragraph_text(self, text: str) -> str:
        """
        三层 fallback 翻译：
        1. 精确匹配（_translation_mapping）
        2. 模糊匹配（normalize 后）
        3. 离线机翻（argostranslate）

        特殊处理：如果段落以中文数字前缀开头（如 "1、"），
        且后半部分是中文章节标题，替换后英文标题本身已含编号，
        需将中文数字前缀转为英文编号格式（如 "1、" → "1. "）。
        """
        if not text or not self._translation_mapping:
            return text

        import re

        # 去除尾随空格（_replace_para_text 不会保留尾随空格，提前清理便于匹配）
        text = text.rstrip()

        # 特殊处理：检测以中文数字前缀开头的段落
        # 分支A：前缀后是中文字符（如 "1、化学品及标识："）
        m_cn = re.match(r'^([0-9１-９一-鿿]+[.、\s]+)([一-鿿])', text)
        # 分支B：前缀后是 ASCII 数字+句点（如 "15、15. Regulatory Information："）
        m_en = re.match(r'^([0-9１-９一-鿿]+[.、\s]+)([0-9]+\.\s*\w)', text)

        for m in (m_cn, m_en):
            if not m:
                continue
            cn_prefix = m.group(1)
            after = text[len(cn_prefix):]
            after_stripped = after.rstrip('：: \t')

            for cn_key, en_val in self._translation_mapping:
                if cn_key == "MATERIAL SAFETY DATA SHEET":
                    continue
                # 精确匹配
                if after_stripped.startswith(cn_key) or after.startswith(cn_key):
                    after_rest = after[len(cn_key):]
                    # after_rest 可能还包含中文内容，需要继续翻译
                    # 先对 after_rest 做 Step 1 全局替换
                    for _cn, _en in self._translation_mapping:
                        after_rest = after_rest.replace(_cn, _en)
                    text = en_val + after_rest
                    return text
                # 去掉尾随标点的 key 匹配（排除纯标点条目如"："）
                cn_key_s = cn_key.rstrip('：: ')
                if cn_key_s and cn_key_s != cn_key and (after_stripped.startswith(cn_key_s) or after.startswith(cn_key_s)):
                    after_rest = after[len(cn_key_s):]
                    for _cn, _en in self._translation_mapping:
                        after_rest = after_rest.replace(_cn, _en)
                    text = en_val + after_rest
                    return text

            # 分支B（ASCII 后缀）匹配了前缀但没找到映射 key → 直接去掉中文前缀
            if m_en:
                text = after
                return text

        # ===== Step 1: 精确匹配（按 key 长度降序全局替换） =====
        original_text = text
        for cn, en in self._translation_mapping:
            text = text.replace(cn, en)

        # 检查是否还有残留中文（基本汉字范围）
        if not re.search(r'[一-鿿㐀-䶿]', text):
            return text  # 已全部翻译

        # ===== Step 2: 模糊匹配（对残留中文段落再次匹配） =====
        normalized = self._normalize(text)
        for cn, en in self._translation_mapping:
            cn_norm = self._normalize(cn)
            if cn_norm != cn and cn_norm in normalized:
                normalized = normalized.replace(cn_norm, en)
        if not re.search(r'[一-鿿㐀-䶿]', normalized):
            return normalized

        # ===== Step 2.5: 中文日期格式转换 =====
        # 将 "YYYY年MM月" 或 "YYYY年M月" 转换为 "YYYY/MM"，避免 argostranslate 误翻日期
        def _convert_chinese_date(t: str) -> str:
            # 匹配中文日期格式：年/月 组合
            t = re.sub(r'(\d{4})年(\d{1,2})月', lambda m: f'{m.group(1)}/{int(m.group(2)):02d}', t)
            return t
        text_with_date = _convert_chinese_date(text)
        if not re.search(r'[一-鿿㐀-䶿]', text_with_date):
            return text_with_date

        # ===== Step 3: argostranslate 机翻 =====
        # 如果文本中已含有英文章节标签（如 "Update Date:"），说明翻译已基本完成，
        # 剩余的中文是待替换的日期占位符，跳过 argostranslate 避免误翻日期格式
        if re.search(r'[A-Za-z]{3,}\s*:', text):
            return text
        # 确保模型已下载（首次调用时下载）
        self._ensure_argos_model()
        return self._translate_with_argos(text)

    def _parse_header(self, doc: Document) -> dict:
        """从文档页眉提取 MSDS编号 和 修订时间"""
        result = {"msds_number": "", "revision_date": ""}
        for section in doc.sections:
            for para in section.header.paragraphs:
                text = para.text
                if "MSDS编号：" in text:
                    # 按"修定时间"分割，取前半部分，再按"MSDS编号："分割
                    if "修定时间" in text:
                        msds_part = text.split("修定时间")[0]
                    elif "修订时间" in text:
                        msds_part = text.split("修订时间")[0]
                    else:
                        msds_part = text
                    if "MSDS编号：" in msds_part:
                        result["msds_number"] = msds_part.split("MSDS编号：")[-1].strip()
                # 修订时间后面紧跟日期 YYYY/MM/DD
                m = re.search(r'(\d{4}/\d{2}/\d{2})', text)
                if m:
                    result["revision_date"] = m.group(1)
        return result

    def parse_msds_file(self, file_path: str) -> dict:
        """
        从旧 MSDS 文件解析出产品信息、成分、理化特性。
        直接从文档解析，不查 customs_codes.json。
        """
        import threading

        # 超时保护：解析文件最多 10 秒，避免卡死
        class TimeoutError(Exception):
            pass

        result = {}
        error = {}

        def _parse():
            try:
                doc = Document(file_path)
                # 提取页眉（MSDS编号 + 修订时间）
                header_info = self._parse_header(doc)
                # 1. 提取产品名称（从段落中找"产品名称："）
                product_name = ""
                for para in doc.paragraphs:
                    if "产品名称：" in para.text:
                        product_name = para.text.split("产品名称：")[-1].strip()
                        break
                # 2. 提取成分表格
                composition = []
                for table in doc.tables:
                    if table.rows:
                        first_row_text = "".join(c.text for c in table.rows[0].cells)
                        if "组分" in first_row_text or "成分" in first_row_text:
                            for row in table.rows[1:]:
                                if len(row.cells) >= 3:
                                    composition.append({
                                        "component_cn": row.cells[0].text.strip(),
                                        "cas": row.cells[1].text.strip(),
                                        "percentage": row.cells[2].text.strip(),
                                    })
                            break
                # 3. 提取理化特性
                import re
                physicochemical = {}
                for para in doc.paragraphs:
                    text = para.text
                    if "外观与性状：" in text:
                        physicochemical["physical_form"] = text.split("外观与性状：")[-1].strip()
                    elif "离子性：" in text:
                        physicochemical["ion_type"] = text.split("离子性：")[-1].strip()
                    elif "离子型：" in text:
                        physicochemical["ion_type"] = text.split("离子型：")[-1].strip()
                    elif re.search(r'[pP][Hh][值]?[:：]\s*([\d±.]+)', text):
                        m = re.search(r'[pP][Hh][值]?[:：]\s*([\d±.]+)', text)
                        if m:
                            physicochemical["ph"] = m.group(1)
                    elif "熔点：" in text:
                        physicochemical["melting_point"] = text.split("：")[-1].strip()
                    elif "沸点" in text and "：" in text:
                        physicochemical["boiling_point"] = text.split("：")[-1].strip()
                    elif "相对密度：" in text:
                        physicochemical["density"] = text.split("相对密度：")[-1].strip()
                    elif "闪点：" in text:
                        physicochemical["flash_point"] = text.split("闪点：")[-1].strip()
                    elif "溶解性：" in text:
                        physicochemical["solubility"] = text.split("溶解性：")[-1].strip()
                result["data"] = {
                    "product_name": product_name,
                    "composition": composition,
                    "physicochemical": physicochemical,
                    "msds_number": header_info["msds_number"],
                    "revision_date": header_info["revision_date"],
                }
            except Exception as e:
                error["e"] = e

        t = threading.Thread(target=_parse)
        t.start()
        t.join(timeout=10)  # 最多等10秒
        if t.is_alive():
            # 线程还在，说明超时了
            raise TimeoutError("解析 MSDS 文件超时（10秒）")
        if error.get("e"):
            raise error["e"]
        if not result:
            raise RuntimeError("解析失败，未知错误")
        return result["data"]

    def search_msds(self, keyword: str) -> list[dict]:
        """
        从 references/MSDS/ 搜索匹配的旧 MSDS 文件。
        返回：[{"name": "xxx.docx", "path": "..."}, ...]
        """
        if not keyword or len(keyword) < 1:
            return []

        keyword_lower = keyword.lower()
        results = []
        msds_path = Path(MSDS_DIR)

        if not msds_path.exists():
            return results

        for f in msds_path.iterdir():
            if f.is_file() and keyword_lower in f.name.lower():
                # 跳过临时文件、test 文件和 PDF 文件（仅支持 docx/doc）
                name_lower = f.name.lower()
                if name_lower.startswith('~$') or name_lower.startswith('test'):
                    continue
                if f.suffix.lower() == '.pdf':
                    continue
                results.append({
                    "name": f.name,
                    "path": str(f)
                })

        return results

    def get_product_data(self, product_name: str) -> Optional[dict]:
        """
        从 customs_codes.json 获取产品数据（按 customs_name 匹配）。
        返回：{"internal_code": "", "customs_name": "", "components": "", "product_appearance": "", ...}
        """
        if not product_name:
            return None

        # 精确匹配 customs_name
        for item in self._customs_codes:
            if item.get("customs_name", "") == product_name:
                return item

        # 模糊匹配
        product_name_lower = product_name.lower()
        for item in self._customs_codes:
            if product_name_lower in item.get("customs_name", "").lower():
                return item

        return None

    def search_products(self, keyword: str) -> list[dict]:
        """
        从 customs_codes.json 搜索产品。
        返回匹配的產品列表（去重）。
        """
        if not keyword or len(keyword) < 1:
            return []

        keyword_lower = keyword.lower()
        results = []
        seen = set()

        for item in self._customs_codes:
            name = item.get("customs_name", "")
            if keyword_lower in name.lower():
                if name not in seen:
                    seen.add(name)
                    results.append({
                        "customs_name": name,
                        "internal_code": item.get("internal_code", ""),
                        "product_appearance": item.get("product_appearance", ""),
                        "components": item.get("components", ""),
                    })

        return results

    def parse_components(self, components_str: str) -> list[dict]:
        """
        解析 components 字段，生成成分表格数据。
        输入示例："柔软吸湿排汗整理剂9016-88-0：90％+水：10％"
        返回：[{"component_cn": "", "cas": "", "percentage": "", "verified": bool}, ...]
        """
        if not components_str:
            return []

        results = []
        # 按 + 拆分多个组分
        parts = components_str.split("+")

        for part in parts:
            part = part.strip()
            if not part:
                continue

            # 按 ：或 : 拆分成分名和含量
            # 格式可能是 "成分名CAS号：含量%" 或 "成分名：含量%"
            # CAS号格式：连续数字加连字符，如 9016-88-0

            # 先尝试提取 CAS 号
            cas_match = re.search(r"(\d{4,7}-\d{2}-\d)", part)
            cas = ""
            component_name = part

            if cas_match:
                cas = cas_match.group(1)
                # 去掉 CAS 号部分，得到成分名
                component_name = part[:cas_match.start()].strip()

            # 提取含量（最后一个出现的百分比）
            pct_match = re.search(r"(\d+(?:\.\d+)?)\s*％", part)
            percentage = ""
            if pct_match:
                percentage = pct_match.group(1) + "%"

            # 验证成分名是否在对照表中
            verified = self._verify_ingredient(component_name)

            results.append({
                "component_cn": component_name,
                "cas": cas,
                "percentage": percentage,
                "verified": verified,
            })

        return results

    def _verify_ingredient(self, ingredient_cn: str) -> bool:
        """验证成分名是否在 ingredient_mapping 中"""
        if not ingredient_cn or not self._ingredient_map:
            return False

        for item in self._ingredient_map:
            if ingredient_cn in item.get("cn_names", []):
                return True
        return False

    def get_cas_by_ingredient_name(self, ingredient_cn: str) -> Optional[str]:
        """通过成分名查找 CAS 号"""
        if not ingredient_cn or not self._ingredient_map:
            return None

        for item in self._ingredient_map:
            if ingredient_cn in item.get("cn_names", []):
                cas_numbers = item.get("cas_numbers", [])
                if cas_numbers:
                    return cas_numbers[0]
        return None

    def translate_ingredient(self, ingredient_cn: str) -> Optional[str]:
        """翻译成分名为英文，支持精确匹配和模糊匹配"""
        if not ingredient_cn:
            return None

        # 1. Exact match
        for item in self._ingredient_map:
            if ingredient_cn in item.get("cn_names", []):
                en_names = item.get("en_names", [])
                if en_names:
                    return en_names[0]

        # 2. Fuzzy match: input is substring of a cn_name (e.g. "甘油" matches "丙三醇（甘油）")
        for item in self._ingredient_map:
            for cn_name in item.get("cn_names", []):
                if ingredient_cn in cn_name:
                    en_names = item.get("en_names", [])
                    if en_names:
                        return en_names[0]

        # 3. Reverse fuzzy: a cn_name is substring of input (e.g. "甘油xx" matches "甘油")
        for item in self._ingredient_map:
            for cn_name in item.get("cn_names", []):
                if cn_name in ingredient_cn:
                    en_names = item.get("en_names", [])
                    if en_names:
                        return en_names[0]

        # 4. Space-normalized match: remove all spaces and compare
        ingredient_no_space = ingredient_cn.replace(" ", "").replace("\u3000", "")
        for item in self._ingredient_map:
            for cn_name in item.get("cn_names", []):
                cn_name_no_space = cn_name.replace(" ", "").replace("\u3000", "")
                if ingredient_no_space == cn_name_no_space:
                    en_names = item.get("en_names", [])
                    if en_names:
                        return en_names[0]
                # Also check substring after removing spaces
                if ingredient_no_space in cn_name_no_space or cn_name_no_space in ingredient_no_space:
                    en_names = item.get("en_names", [])
                    if en_names:
                        return en_names[0]

        return None

    def translate_appearance(self, appearance_cn: str) -> Optional[str]:
        """翻译外观描述为英文"""
        return self._appearance_map.get(appearance_cn)

    def translate_product_name(self, product_name_cn: str) -> Optional[str]:
        """翻译产品名为英文"""
        return self._products_name_map.get(product_name_cn)

    def translate_ion_type(self, ion_type_cn: str) -> Optional[str]:
        """翻译离子性为英文"""
        return self._ion_type_map.get(ion_type_cn)

    def translate_text_fallback(self, text_cn: str) -> str:
        """
        机翻中文到英文的 fallback 方法。
        优先使用本地映射，其次使用简单的字符替换翻译。
        """
        if not text_cn:
            return text_cn

        # 先尝试精确匹配 appearance
        appearance_en = self._appearance_map.get(text_cn)
        if appearance_en:
            return appearance_en

        # 尝试产品名映射
        product_en = self._products_name_map.get(text_cn)
        if product_en:
            return product_en

        # 尝试离子性映射
        ion_en = self._ion_type_map.get(text_cn)
        if ion_en:
            return ion_en

        # 尝试成分映射
        ingredient_en = self.translate_ingredient(text_cn)
        if ingredient_en:
            return ingredient_en

        # 无法翻译，返回原文并标注
        return text_cn

    def generate_msds(
        self,
        msds_file_path: str,
        product_data: dict,
        edits: dict
    ) -> tuple[bytes, str]:
        """
        基于旧 MSDS 文件，应用编辑内容，生成新 MSDS docx。

        Args:
            msds_file_path: 旧 MSDS 文件路径
            product_data: 从 customs_codes.json 获取的产品数据
            edits: 用户编辑的内容 {
                "product_name": str,
                "composition": [{"component_cn": "", "cas": "", "percentage": ""}, ...],
                "physicochemical": {"外观与性状": "", "离子性": "", ...}
            }

        Returns:
            (docx_bytes, doc_key)
        """
        # 读取旧 MSDS 文件
        doc = Document(msds_file_path)

        # 1. 修改产品名称（如果编辑了）
        if edits.get("product_name"):
            self._update_product_name(doc, edits["product_name"])

        # 2. 修改成分表格
        if edits.get("composition"):
            self._update_composition_table(doc, edits["composition"])

        # 3. 修改理化特性
        if edits.get("physicochemical"):
            self._update_physicochemical(doc, edits["physicochemical"])

        # 4. 修改页眉（MSDS编号 + 修订时间）
        if edits.get("msds_number") or edits.get("revision_date"):
            self._update_header(doc, edits.get("msds_number", ""), edits.get("revision_date", ""))

        # 5. 更新第16条"其他资讯"中的版次和更新日期
        if edits.get("revision_date"):
            self._update_section16(doc, edits["revision_date"])

        # 保存到 BytesIO
        buf = BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        # 生成 doc_key
        import time
        product_name = edits.get("product_name", product_data.get("customs_name", "unknown"))
        doc_key = f"msds_{product_name}_{int(time.time())}"

        return content, doc_key

    def _update_product_name(self, doc: Document, new_name: str):
        """更新文档中的产品名称。保留原始段落格式和字体属性。"""
        from docx.shared import Pt, Pt as PtType
        from copy import copy

        for para in doc.paragraphs:
            if "产品名称：" in para.text:
                full_text = para.text
                idx = full_text.find("产品名称：")
                prefix = full_text[:idx]

                # 保存原始段落格式（line_spacing 等）
                pf = para.paragraph_format
                saved_line_spacing = pf.line_spacing
                saved_first_indent = pf.first_line_indent
                saved_left_indent = pf.left_indent

                # 保存原始 run 的字体属性（取第一个 run 的字体）
                first_run = para.runs[0] if para.runs else None
                orig_font_name = first_run.font.name if first_run else None
                orig_font_size = first_run.font.size if first_run else None

                # 清空所有 run 并重建
                for run in para.runs:
                    run.text = ""
                para.clear()

                new_run = para.add_run(prefix + f"产品名称：{new_name}")
                # 保持与原始一致的字体（继承而非显式设置）
                if orig_font_name:
                    new_run.font.name = orig_font_name
                elif orig_font_size:
                    # 原始有字号但无显式字体名，保持字号继承
                    new_run.font.size = orig_font_size

                # 恢复段落格式
                pf.line_spacing = saved_line_spacing
                pf.first_line_indent = saved_first_indent
                pf.left_indent = saved_left_indent
                break

    def _update_composition_table(self, doc: Document, composition: list):
        """更新成分表格，保留原有字体。"""
        from docx.shared import Pt

        tables = doc.tables
        if not tables:
            return

        # 找到成分表格（第一行包含"组分"或"成分"）
        comp_table = None
        for table in tables:
            if table.rows:
                first_row_text = "".join(c.text for c in table.rows[0].cells)
                if "组分" in first_row_text or "成分" in first_row_text:
                    comp_table = table
                    break

        if not comp_table:
            return

        # 获取原始单元格的字体样式（从第一行数据行获取）
        sample_row = comp_table.rows[1] if len(comp_table.rows) > 1 else None
        sample_fonts = []
        if sample_row and len(sample_row.cells) >= 3:
            for cell in sample_row.cells:
                if cell.paragraphs and cell.paragraphs[0].runs:
                    first_run = cell.paragraphs[0].runs[0]
                    sample_fonts.append({
                        "name": first_run.font.name,
                        "size": first_run.font.size,
                    })
                else:
                    sample_fonts.append({"name": None, "size": None})

        # 更新每一行
        for i, item in enumerate(composition):
            row_idx = i + 1  # 跳过表头
            if row_idx >= len(comp_table.rows):
                # 现有行不够，新增一行（复制最后一行的结构）
                new_row = comp_table.add_row()
                # 复制最后一行的单元格样式
                if comp_table.rows and len(comp_table.rows) > 1:
                    last_row = comp_table.rows[-2]  # 最后一 个有数据的行
                    for col_idx in range(min(len(new_row.cells), len(last_row.cells))):
                        src_cell = last_row.cells[col_idx]
                        dst_cell = new_row.cells[col_idx]
                        # 复制段落格式（paragraph_format 是只读属性，只能逐属性复制）
                        if src_cell.paragraphs and dst_cell.paragraphs:
                            src_para = src_cell.paragraphs[0]
                            dst_para = dst_cell.paragraphs[0] if dst_cell.paragraphs else dst_cell.add_paragraph()
                            src_fmt = src_para.paragraph_format
                            dst_fmt = dst_para.paragraph_format
                            # 逐属性复制（ParagraphFormat 的属性都是可写的）
                            dst_fmt.alignment = src_fmt.alignment
                            dst_fmt.line_spacing = src_fmt.line_spacing
                            dst_fmt.line_spacing_rule = src_fmt.line_spacing_rule
                            dst_fmt.space_before = src_fmt.space_before
                            dst_fmt.space_after = src_fmt.space_after
                            dst_fmt.left_indent = src_fmt.left_indent
                            dst_fmt.right_indent = src_fmt.right_indent
                            dst_fmt.first_line_indent = src_fmt.first_line_indent
                            # 复制 run 样式（如果存在）
                            if src_para.runs:
                                src_run = src_para.runs[0]
                                dst_run = dst_para.add_run("")
                                dst_run.font.name = src_run.font.name
                                dst_run.font.size = src_run.font.size
                                dst_run.font.bold = src_run.font.bold
                                if src_run.font.color.rgb:
                                    dst_run.font.color.rgb = src_run.font.color.rgb
                row = new_row
            else:
                row = comp_table.rows[row_idx]
            if len(row.cells) >= 3:
                new_values = [
                    item.get("component_cn", ""),
                    item.get("cas", ""),
                    item.get("percentage", ""),
                ]
                for col_idx, new_val in enumerate(new_values):
                    if col_idx >= len(row.cells):
                        continue
                    cell = row.cells[col_idx]
                    # 直接修改现有 run 的文本，保留字体
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        # 只修改第一个 run 的文本，清空其余 run
                        para = cell.paragraphs[0]
                        for r_idx, run in enumerate(para.runs):
                            if r_idx == 0:
                                run.text = new_val
                            else:
                                run.text = ""
                        # 如果没有 run，创建一个
                        if not para.runs:
                            new_run = para.add_run(new_val)
                            if col_idx < len(sample_fonts) and sample_fonts[col_idx]["name"]:
                                new_run.font.name = sample_fonts[col_idx]["name"]
                            if col_idx < len(sample_fonts) and sample_fonts[col_idx]["size"]:
                                new_run.font.size = sample_fonts[col_idx]["size"]
                    else:
                        # 没有段落，创建一个
                        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        new_run = para.add_run(new_val)
                        if col_idx < len(sample_fonts) and sample_fonts[col_idx]["name"]:
                            new_run.font.name = sample_fonts[col_idx]["name"]
                        if col_idx < len(sample_fonts) and sample_fonts[col_idx]["size"]:
                            new_run.font.size = sample_fonts[col_idx]["size"]

    def _update_header(self, doc: Document, msds_number: str, revision_date: str):
        """更新文档页眉中的 MSDS编号 和 修订时间"""
        for section in doc.sections:
            for para in section.header.paragraphs:
                text = para.text
                new_text = text

                # 替换 MSDS编号（用 split 避免正则贪婪匹配后面的"修订时间"文字）
                if msds_number and "MSDS编号：" in text:
                    # 找到"修订时间"或"修定时间"的位置
                    marker_name = ""
                    marker_pos = -1
                    for m_name in ["修定时间", "修订时间", "修订日期"]:
                        pos = text.find(m_name)
                        if pos >= 0:
                            marker_pos = pos
                            marker_name = m_name
                            break
                    if marker_pos > 0:
                        # msds_prefix = "MSDS编号："（到冒号之后）
                        msds_prefix = text.split("MSDS编号：")[0] + "MSDS编号："
                        # 提取原始日期
                        m = re.search(r'(\d{4}/\d{2}/\d{2})', text)
                        orig_date = m.group(1) if m else ""
                        # 用新编号 + 日期重建
                        date_to_use = revision_date if revision_date else orig_date
                        new_text = f"{msds_prefix}{msds_number}   {marker_name}：{date_to_use}"
                    else:
                        # 没有找到时间标记，直接替换MSDS编号部分
                        msds_prefix = text.split("MSDS编号：")[0] + "MSDS编号："
                        new_text = msds_prefix + msds_number

                # 替换修订时间（当 revision_date 有值时覆盖原有日期）
                if revision_date:
                    m = re.search(r'(\d{4}/\d{2}/\d{2})', new_text)
                    if m:
                        new_text = re.sub(r'(\d{4}/\d{2}/\d{2})', revision_date, new_text, count=1)

                # 只有文本变化时才更新
                if new_text != text:
                    # 保存原始段落格式
                    pf = para.paragraph_format
                    saved_line_spacing = pf.line_spacing
                    saved_first_indent = pf.first_line_indent
                    saved_left_indent = pf.left_indent

                    # 保存原始 run 的字体属性
                    first_run = para.runs[0] if para.runs else None
                    orig_font_name = first_run.font.name if first_run else None
                    orig_font_size = first_run.font.size if first_run else None

                    # 检测段落是否包含绘图（如 logo/印章）
                    from lxml import etree
                    para_xml = para._element.xml
                    has_drawing = '<w:drawing' in para_xml

                    if has_drawing:
                        # 含绘图的段落：只清空文字run，保留绘图run
                        wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                        for r in para._element.findall(f'{{{wns}}}r'):
                            has_r_drawing = False
                            if r.find(f'{{{wns}}}drawing') is not None:
                                has_r_drawing = True
                            else:
                                r_xml = etree.tostring(r, encoding='unicode')
                                if '<w:drawing' in r_xml or '<w:anchor' in r_xml:
                                    has_r_drawing = True
                            if not has_r_drawing:
                                for t in r.findall(f'{{{wns}}}t'):
                                    t.text = ''
                        new_run = para.add_run(new_text)
                        if orig_font_name:
                            new_run.font.name = orig_font_name
                        elif orig_font_size:
                            new_run.font.size = orig_font_size
                    else:
                        # 普通段落：完全替换
                        for run in para.runs:
                            run.text = ""
                        para.clear()
                        new_run = para.add_run(new_text)
                        if orig_font_name:
                            new_run.font.name = orig_font_name
                        elif orig_font_size:
                            new_run.font.size = orig_font_size

                    # 恢复段落格式
                    pf.line_spacing = saved_line_spacing
                    pf.first_line_indent = saved_first_indent
                    pf.left_indent = saved_left_indent

    def _update_section16(self, doc: Document, revision_date: str):
        """
        更新第16条"其他资讯"中的版次和更新日期。
        revision_date 格式: YYYY/MM/DD
        - 版次: YYYY-MM（如 2026-05）
        - 更新日期: YYYY年M月（如 2026年5月）
        """
        # 解析日期
        parts = revision_date.split("/")
        if len(parts) < 2:
            return
        year = parts[0]
        month = parts[1]
        # 去掉前导零
        month_str = str(int(month))
        banci = f"{year}-{month}"           # 版次: 2026-05
        gengxin = f"{year}年{month_str}月"  # 更新日期: 2026年5月

        for para in doc.paragraphs:
            text = para.text
            if "版次" not in text:
                continue

            new_text = text
            # 替换版次: 找到 "版次：" 后的 YYYY-MM
            m_ban = re.search(r'(版次[：:]\s*)([\d]{4}-[\d]{2})', new_text)
            if m_ban:
                new_text = new_text[:m_ban.start()] + m_ban.group(1) + banci + new_text[m_ban.end():]

            # 替换更新日期: 找到 "更新日期：" 后的日期
            m_geng = re.search(r'(更新日期[：:]\s*)([\d]+年[\d]+月)', new_text)
            if m_geng:
                new_text = new_text[:m_geng.start()] + m_geng.group(1) + gengxin + new_text[m_geng.end():]

            if new_text != text:
                # 保存并恢复段落格式
                pf = para.paragraph_format
                saved_line_spacing = pf.line_spacing
                saved_first_indent = pf.first_line_indent
                saved_left_indent = pf.left_indent
                first_run = para.runs[0] if para.runs else None
                orig_font_name = first_run.font.name if first_run else None
                orig_font_size = first_run.font.size if first_run else None

                for run in para.runs:
                    run.text = ""
                para.clear()
                new_run = para.add_run(new_text)
                if orig_font_name:
                    new_run.font.name = orig_font_name
                elif orig_font_size:
                    new_run.font.size = orig_font_size

                pf.line_spacing = saved_line_spacing
                pf.first_line_indent = saved_first_indent
                pf.left_indent = saved_left_indent
            break

    # ===================== English MSDS Template Methods =====================

    def generate_msds_english_from_template(
        self,
        msds_file_path: str,
        parsed_data: dict,
        edits: dict
    ) -> tuple[bytes, str]:
        """
        基于英文 MSDS 模板 + 用户编辑字段，生成英文 MSDS docx。

        Args:
            msds_file_path: 中文 MSDS 文件路径（用于解析原始数据，但不使用其内容）
            parsed_data: parse_msds_file() 返回的中文 MSDS 数据
            edits: 用户编辑的内容 {
                "product_name": str,
                "composition": [{"component_cn": "", "cas": "", "percentage": ""}, ...],
                "physicochemical": {"appearance": "", "ion_type": "", "ph": "", ...},
                "msds_number": str,
                "revision_date": str,
                "language": "cn" | "en"
            }

        Returns:
            (docx_bytes, doc_key)
        """
        import shutil
        import time

        # 1. 复制英文模板到临时文件（模板只读）
        template_path = ENGLISH_MSDS_TEMPLATE
        if not Path(template_path).exists():
            raise FileNotFoundError(f"英文 MSDS 模板不存在: {template_path}")

        # 读取模板生成文档
        doc = Document(template_path)

        # 2. 翻译产品名
        product_name_en = self._translate_product_name_en(
            edits.get("product_name") or parsed_data.get("product_name", "")
        )

        # 3. 翻译成分（中文→英文）
        composition = edits.get("composition") or parsed_data.get("composition", [])
        composition_en = []
        for item in composition:
            comp_cn = item.get("component_cn", "")
            comp_en = self._translate_text_en(comp_cn)
            composition_en.append({
                "component_en": comp_en,
                "cas": item.get("cas", ""),
                "percentage": item.get("percentage", ""),
            })

        # 4. 填充各字段
        self._fill_english_product_name(doc, product_name_en)
        self._fill_english_composition_table(doc, composition_en)
        self._fill_english_physicochemical(doc, edits.get("physicochemical", {}))
        self._fill_english_header(
            doc,
            edits.get("msds_number", ""),
            edits.get("revision_date", "")
        )
        self._fill_english_section16(
            doc,
            edits.get("msds_number", ""),
            edits.get("revision_date", "")
        )

        # 5. 保存到 BytesIO
        buf = BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        # 生成 doc_key
        doc_key = f"msds_en_{product_name_en}_{int(time.time())}"
        return content, doc_key

    def _fill_english_product_name(self, doc: Document, product_name_en: str):
        """
        替换文档中的产品名称。
        两个目标位置：
        1. 标题下方的独立产品名段落（如 "FIXING AGENT"）
        2. Section 1 的 "Product Name: XXX" 段落
        """
        if not product_name_en:
            return

        for para in doc.paragraphs:
            text = para.text
            if not text:
                continue

            # 位置1：独立产品名段落（如 "FIXING AGENT"）
            # 排除所有含冒号的行（含 Product Name / 字段名 等）
            has_colon = ':' in text or '：' in text  # 半角或全角冒号都排除
            if not has_colon:
                stripped = text.strip()
                # 排除标题、章节开头、固定公司信息行、常见英文 MSDS 段落开头
                skip_prefixes = [
                    "MATERIAL SAFETY DATA SHEET",
                    "Company", "Tel", "Fax", "E-mail", "Website", "Address",
                    "Component", "CAS NO", "W/W",
                    "Emergency", "Risk", "Skin", "Eye", "Inhalation", "Ingestion",
                    "Suggestion", "Damage", "Hazard", "Precautionary", "Burning",
                    "Health", "Acute", "Environmental", "Stability",
                    "Toxicological", "Ecological", "Disposal",
                    "Transport", "Regulatory", "Revision", "Update",
                    "Other", "According", "Department",
                    "Personal Protective", "Forbid", "Prevent", "If it can",
                    "Law on", "Regulations", "It is forbidden", "IMO",
                    "1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.",
                    "10.", "11.", "12.", "13.", "14.", "15.", "16.",
                ]
                # 产品名段落特征：全大写（如 FIXING AGENT），排除普通句子
                if (stripped and len(stripped) > 2 and len(stripped) < 80
                        and stripped.isupper()
                        and not any(stripped.startswith(p) for p in skip_prefixes)):
                    self._replace_para_text(para, product_name_en)
                    continue

            # 位置2：Section 1 的 "Product Name:XXX"（含全角冒号）
            if 'Product Name' in text and ('_' in text or ':' in text or '：' in text):
                # 用全角/半角冒号分割，取 prefix
                for sep in ['：', ':']:
                    if sep in text:
                        colon_idx = text.index(sep)
                        prefix = text[:colon_idx + 1]
                        after = text[colon_idx + 1:]
                        if after.strip() and after.strip() != product_name_en:
                            self._replace_para_text(para, prefix + " " + product_name_en)
                        break

    def _fill_english_composition_table(self, doc: Document, composition_en: list):
        """更新英文模板的成分表格（Section 3）"""
        tables = doc.tables
        if not tables:
            return

        # 英文模板只有1个表格（Section 3 - Main Component）
        table = tables[0]

        # 表头已是英文: Component | CAS NO. | W/W %
        # 确保表头正确
        if len(table.rows) > 0:
            hdr_cells = table.rows[0].cells
            expected_headers = ["Component", "CAS NO.", "W/W %"]
            for i, cell in enumerate(hdr_cells[:3]):
                cell_text = cell.text.strip()
                # 表头可能因空格/大小写略有不同，尝试模糊匹配
                matched = False
                for exp in expected_headers:
                    if exp.lower() in cell_text.lower() or cell_text.lower() in exp.lower():
                        if cell.paragraphs and cell.paragraphs[0].runs:
                            for r in cell.paragraphs[0].runs:
                                if r.text.strip():
                                    r.text = exp
                                    matched = True
                                    break
                        break
                if not matched and i < len(expected_headers):
                    # 完全没有匹配，直接覆盖文本
                    if cell.paragraphs:
                        cell.paragraphs[0].clear()
                        cell.paragraphs[0].add_run(expected_headers[i])

        # 获取原始单元格字体样式
        sample_fonts = []
        if len(table.rows) > 1:
            sample_row = table.rows[1]
            for cell in sample_row.cells[:3]:
                if cell.paragraphs and cell.paragraphs[0].runs:
                    first_run = cell.paragraphs[0].runs[0]
                    sample_fonts.append({
                        "name": first_run.font.name,
                        "size": first_run.font.size,
                    })
                else:
                    sample_fonts.append({"name": None, "size": None})
        else:
            sample_fonts = [{"name": None, "size": None}] * 3

        # 调整行数：用户有 N 行，模板至少有2行（表头+1行数据）
        # 如果不够，复制最后一行新增
        while len(table.rows) - 1 < len(composition_en):
            # 复制最后一行
            src_row = table.rows[-1]
            new_row = table.add_row()
            # 复制单元格样式
            for col_idx in range(min(len(new_row.cells), len(src_row.cells))):
                src_cell = src_row.cells[col_idx]
                dst_cell = new_row.cells[col_idx]
                if src_cell.paragraphs and dst_cell.paragraphs:
                    src_para = src_cell.paragraphs[0]
                    dst_para = dst_cell.paragraphs[0] if dst_cell.paragraphs else dst_cell.add_paragraph()
                    src_fmt = src_para.paragraph_format
                    dst_fmt = dst_para.paragraph_format
                    dst_fmt.alignment = src_fmt.alignment
                    dst_fmt.line_spacing = src_fmt.line_spacing
                    dst_fmt.line_spacing_rule = src_fmt.line_spacing_rule
                    dst_fmt.space_before = src_fmt.space_before
                    dst_fmt.space_after = src_fmt.space_after
                    dst_fmt.left_indent = src_fmt.left_indent
                    dst_fmt.right_indent = src_fmt.right_indent
                    dst_fmt.first_line_indent = src_fmt.first_line_indent
                    if src_para.runs:
                        src_run = src_para.runs[0]
                        dst_run = dst_para.add_run("")
                        dst_run.font.name = src_run.font.name
                        dst_run.font.size = src_run.font.size
                        dst_run.font.bold = src_run.font.bold
                        if src_run.font.color.rgb:
                            dst_run.font.color.rgb = src_run.font.color.rgb

        # 如果用户行数少于模板行数，删除多余数据行（保留表头）
        while len(table.rows) - 1 > len(composition_en) and len(table.rows) > 2:
            tbl = table._tbl
            tbl.remove(table.rows[-1]._tr)

        # 填充数据
        for i, item in enumerate(composition_en):
            row_idx = i + 1
            if row_idx >= len(table.rows):
                break
            row = table.rows[row_idx]
            if len(row.cells) >= 3:
                new_values = [
                    item.get("component_en", ""),
                    item.get("cas", ""),
                    item.get("percentage", ""),
                ]
                for col_idx, new_val in enumerate(new_values):
                    if col_idx >= len(row.cells):
                        continue
                    cell = row.cells[col_idx]
                    if cell.paragraphs and cell.paragraphs[0].runs:
                        para = cell.paragraphs[0]
                        for r_idx, run in enumerate(para.runs):
                            if r_idx == 0:
                                run.text = new_val
                            else:
                                run.text = ""
                    else:
                        para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        new_run = para.add_run(new_val)
                        if sample_fonts[col_idx]["name"]:
                            new_run.font.name = sample_fonts[col_idx]["name"]
                        if sample_fonts[col_idx]["size"]:
                            new_run.font.size = sample_fonts[col_idx]["size"]

    def _fill_english_physicochemical(self, doc: Document, physicochemical: dict):
        """更新英文模板的理化特性（Section 9）"""
        if not physicochemical:
            return

        # 英文模板字段名 → (表单字段名, 翻译函数)
        field_map = {
            "Appearance:": ("appearance", self.translate_appearance),
            "Ionicity:": ("ion_type", self.translate_ion_type),
            "pH:": ("ph", None),
            "Melting point:": ("melting_point", None),
            "Boiling point:": ("boiling_point", None),
            "Boiling point(°C):": ("boiling_point", None),
            "Density(°C):": ("density", None),
            "Relative density:": ("density", None),
            "Flash Point:": ("flash_point", None),
            "Flame point:": ("flash_point", None),
            "Solubility:": ("solubility", self._translate_text_en),
        }

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 匹配字段名
            matched_field = None
            for en_field in field_map:
                if text.startswith(en_field) or en_field in text:
                    matched_field = en_field
                    break
            if not matched_field:
                continue

            form_key, translator = field_map[matched_field]
            raw_value = physicochemical.get(form_key, "")
            if not raw_value:
                # 用户未填写，保留原模板值不变
                continue

            # 翻译值：先尝试专用翻译函数，再用通用翻译 fallback
            display_value = raw_value
            if translator:
                translated = translator(raw_value)
                if translated and translated != raw_value:
                    display_value = translated
                else:
                    # 专用翻译无结果，用通用 fallback 再试一次
                    fallback = self._translate_text_en(raw_value)
                    if fallback and fallback != raw_value:
                        display_value = fallback
            else:
                # 无专用翻译函数，直接用通用翻译（数值类字段如 pH / 熔点 / 沸点 等）
                fallback = self._translate_text_en(raw_value)
                if fallback and fallback != raw_value:
                    display_value = fallback

            # 替换冒号后的值
            colon_pos = text.find(":")
            if colon_pos < 0:
                continue

            # 构建新文本：字段名 + " " + 值
            new_text = matched_field + " " + display_value
            # 保留末尾标点（.,）
            after_colon = text[colon_pos + 1:]
            trailing = after_colon.strip()
            if trailing and trailing[-1] in ('.', ',', ';'):
                new_text = new_text + trailing[-1]
            elif after_colon.strip().endswith(";"):
                trailing_punct = after_colon.strip()[-1]
                if not new_text.endswith(trailing_punct):
                    new_text = new_text + trailing_punct

            if new_text != text:
                self._replace_para_text(para, new_text)

    def _fill_english_header(self, doc: Document, msds_number: str, revision_date: str):
        """更新页眉中的 MSDS No. 和 Revision time"""
        for section in doc.sections:
            for para in section.header.paragraphs:
                text = para.text
                if not text:
                    continue
                new_text = text

                # 模板格式: "MSDS No.：HHJS-2633     Revision time：2026/01/01"
                # 全角冒号 "：" 分割
                msds_label = "MSDS No.："
                rev_label = "Revision time："

                # 替换 MSDS No. 值
                if msds_label in new_text:
                    msds_idx = new_text.index(msds_label) + len(msds_label)
                    rev_start = new_text.index(rev_label)
                    old_msds_val = new_text[msds_idx:rev_start].strip()
                    new_text = new_text.replace(msds_label + old_msds_val, msds_label + msds_number, 1)

                # 替换 Revision time 值
                if rev_label in new_text:
                    old_rev_val = new_text[new_text.index(rev_label) + len(rev_label):].strip()
                    new_text = new_text.replace(rev_label + old_rev_val, rev_label + revision_date, 1)

                if new_text != text:
                    self._replace_para_text(para, new_text)

    def _fill_english_section16(self, doc: Document, msds_number: str, revision_date: str):
        """更新 Section 16 的 Revision 和 Update Date"""
        for para in doc.paragraphs:
            text = para.text
            new_text = text

            # 模板格式: "Revision: JAN.,2026             Update Date:JAN,2026"
            # 1. 替换 Revision: 后的旧值（如 "JAN.,2026"）
            if "Revision:" in new_text:
                rev_label = "Revision: "
                if rev_label in new_text:
                    rev_idx = new_text.index(rev_label) + len(rev_label)
                    next_label_idx = new_text.find("Update Date:", rev_idx)
                    if next_label_idx >= 0:
                        old_rev_val = new_text[rev_idx:next_label_idx].strip()
                        new_text = new_text.replace(rev_label + old_rev_val, rev_label + msds_number, 1)

            # 2. 替换 Update Date: 后的旧值（如 "JAN,2026"）
            if "Update Date:" in new_text:
                date_label = "Update Date:"
                if date_label in new_text:
                    date_idx = new_text.index(date_label) + len(date_label)
                    old_date_val = new_text[date_idx:].strip()
                    new_text = new_text.replace(date_label + old_date_val, date_label + revision_date, 1)

            if new_text != text:
                self._replace_para_text(para, new_text)

    # ===================== English MSDS Methods (old translate-based) =====================

    def _get_english_reference_file(self) -> Optional[str]:
        """获取英文 MSDS 参考文件路径"""
        en_dir = Path(ENGLISH_MSDS_DIR)
        if not en_dir.exists():
            return None
        # 优先找带"英文"或"EN"的 docx 文件
        for f in en_dir.iterdir():
            if f.is_file() and f.suffix.lower() in [".docx", ".doc"]:
                name_lower = f.name.lower()
                if "英文" in name_lower or "english" in name_lower or "_en" in name_lower or f.name.startswith("MSDS for"):
                    return str(f)
        # fallback: 找任意 docx
        for f in en_dir.iterdir():
            if f.is_file() and f.suffix.lower() == ".docx":
                return str(f)
        return None

    def extract_english_msds_structure(self, file_path: str = None) -> dict:
        """
        从英文 MSDS 参考文件提取英文 MSDS 的标准结构。
        缓存结果避免重复解析。
        """
        cache_key = file_path or "default"
        if cache_key in self._english_structure_cache:
            return self._english_structure_cache[cache_key]

        if not file_path:
            file_path = self._get_english_reference_file()
        if not file_path or not Path(file_path).exists():
            # 返回默认结构
            return self._get_default_english_structure()

        try:
            doc = Document(file_path)
        except Exception:
            return self._get_default_english_structure()

        structure = {
            "section_titles": [],
            "fixed_texts": [],
            "composition_headers": COMPOSITION_HEADERS_EN,
            "physicochemical_keys": list(PHYSLICOCHEMICAL_KEYS_EN.keys()),
            "section16_format": {
                "banci_label": "Revision No.",
                "gengxin_label": "Revision Date.",
            },
        }

        # 提取 section 标题（从段落中识别 1. 2. 之类的）
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # 匹配类似 "1. Identification" 或 "MATERIAL SAFETY DATA SHEET"
            if re.match(r'^\d+\.\s+\w', text) or text == "MATERIAL SAFETY DATA SHEET":
                structure["section_titles"].append(text)

        # 提取理化特性 key（从英文参考文件中识别）
        physico_keys_found = {}
        for para in doc.paragraphs:
            text = para.text.strip()
            for field, markers in PHYSLICOCHEMICAL_KEYS_EN.items():
                for marker in markers:
                    if marker in text:
                        # 提取冒号后的值
                        if "：" in text:
                            val = text.split("：")[-1].strip()
                        elif ": " in text:
                            val = text.split(": ")[-1].strip()
                        else:
                            val = ""
                        physico_keys_found[field] = val
                        break

        structure["physicochemical_keys"] = physico_keys_found

        self._english_structure_cache[cache_key] = structure
        return structure

    def _get_default_english_structure(self) -> dict:
        """返回默认的英文 MSDS 结构"""
        return {
            "section_titles": SECTION_TITLES_EN,
            "fixed_texts": [
                "Company:",
                "Tel:",
                "Fax:",
                "Section 1: Identification",
                "Section 2: Hazards",
                "Section 3: Composition",
                "Section 4: First Aid",
                "Section 5: Fire Fighting",
                "Section 6: Accidental Release",
                "Section 7: Handling and Storage",
                "Section 8: Exposure Controls",
                "Section 9: Physical and Chemical Properties",
                "Section 10: Stability",
                "Section 11: Toxicology",
                "Section 12: Ecology",
                "Section 13: Disposal",
                "Section 14: Transport",
                "Section 15: Regulation",
                "Section 16: Other Information",
            ],
            "composition_headers": COMPOSITION_HEADERS_EN,
            "physicochemical_keys": list(PHYSLICOCHEMICAL_KEYS_EN.keys()),
            "section16_format": {
                "banci_label": "Revision No.",
                "gengxin_label": "Revision Date.",
            },
        }

    def generate_msds_english(
        self,
        msds_file_path: str,
        parsed_data: dict,
        edits: dict
    ) -> tuple[bytes, str]:
        """
        基于中文 MSDS 文件，生成英文 MSDS docx。
        以中文 MSDS 文件为底版，保留所有样式，只替换文字内容为英文。

        Args:
            msds_file_path: 中文 MSDS 文件路径
            parsed_data: parse_msds_file() 返回的中文 MSDS 数据
            edits: 用户编辑的内容

        Returns:
            (docx_bytes, doc_key)
        """
        # 直接打开用户选中的中文 MSDS 文件作为底版（保留所有样式）
        doc = Document(msds_file_path)

        # 翻译产品名
        product_name_en = self._translate_product_name_en(
            edits.get("product_name") or parsed_data.get("product_name", "")
        )

        # 翻译成分
        composition = edits.get("composition") or parsed_data.get("composition", [])
        composition_en = []
        for item in composition:
            comp_cn = item.get("component_cn", "")
            comp_en = self._translate_text_en(comp_cn)
            composition_en.append({
                "component_en": comp_en,
                "cas": item.get("cas", ""),
                "percentage": item.get("percentage", ""),
            })

        # 第一步：通用段落翻译（字段名/章节标题/固定短语/全角标点全部替换）
        for para in doc.paragraphs:
            if para.text.strip():
                translated = self._translate_paragraph_text(para.text)
                self._replace_para_text(para, translated)

        # 第二步：翻译表格单元格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            translated = self._translate_paragraph_text(para.text)
                            self._replace_para_text(para, translated)

        # 第三步：更新成分表格（用英文数据替换表头和数据）
        self._update_english_composition_table(doc, composition_en)

        # 第四步：更新页眉（MSDS编号和修订时间）
        self._update_english_header(doc, edits.get("msds_number", ""), edits.get("revision_date", ""))

        # 第五步：更新 Section 16 的版次和更新日期数值
        revision_date = edits.get("revision_date", "")
        if revision_date:
            self._update_section16_values(doc, revision_date)

        # 第六步：翻译理化特性字段值（Appearance、Ionicity、Solubility 等）
        self._translate_physicochemical_values(doc)

        # 第七步：翻译产品名称字段值
        if product_name_en:
            self._translate_product_name_field(doc, product_name_en)

        # 保存到 BytesIO
        buf = BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        # 生成 doc_key
        import time
        doc_key = f"msds_en_{product_name_en}_{int(time.time())}"

        return content, doc_key

    def _translate_product_name_en(self, product_name_cn: str) -> str:
        """翻译产品名为英文"""
        if not product_name_cn:
            return ""
        en = self.translate_product_name(product_name_cn)
        if en:
            return en
        return self.translate_text_fallback(product_name_cn)

    def _replace_para_text(self, para, new_text: str):
        """
        替换段落文本（保留段落格式和字体）。
        如果段落包含绘图（如背景logo/印章），只替换文字部分，保留绘图。
        """
        from lxml import etree

        pf = para.paragraph_format
        saved_line_spacing = pf.line_spacing
        saved_first_indent = pf.first_line_indent
        saved_left_indent = pf.left_indent

        # 检测段落是否包含绘图（<w:drawing> 元素）
        para_xml = para._element.xml
        has_drawing = '<w:drawing' in para_xml

        first_run = para.runs[0] if para.runs else None
        orig_font_name = first_run.font.name if first_run else None
        orig_font_size = first_run.font.size if first_run else None

        if has_drawing:
            # 段落含绘图（如背景logo/印章）：
            # 清空所有文字run的文字，保留绘图run，然后在末尾添加新文字run
            wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            for r in para._element.findall(f'{{{wns}}}r'):
                # 检查run是否含<w:drawing>（直接在r下或作为子元素）
                has_r_drawing = False
                if r.find(f'{{{wns}}}drawing') is not None:
                    has_r_drawing = True
                else:
                    # 用字符串方式检测子元素中是否有drawing
                    r_xml = etree.tostring(r, encoding='unicode')
                    if '<w:drawing' in r_xml or '<w:anchor' in r_xml:
                        has_r_drawing = True
                if not has_r_drawing:
                    # 文字run：清空文字
                    for t in r.findall(f'{{{wns}}}t'):
                        t.text = ''
                # else: 绘图run，保留
            # 添加新的文字run
            new_run = para.add_run(new_text)
            if orig_font_name:
                new_run.font.name = orig_font_name
            elif orig_font_size:
                new_run.font.size = orig_font_size
        else:
            # 普通段落：完全替换
            for run in para.runs:
                run.text = ""
            para.clear()
            new_run = para.add_run(new_text)
            if orig_font_name:
                new_run.font.name = orig_font_name
            elif orig_font_size:
                new_run.font.size = orig_font_size

        pf.line_spacing = saved_line_spacing
        pf.first_line_indent = saved_first_indent
        pf.left_indent = saved_left_indent

    def _translate_text_en(self, text_cn: str) -> str:
        """
        翻译任意中文文本为英文。
        优先使用映射表，找不到则调用 argostranslate 离线翻译。
        """
        if not text_cn:
            return ""
        # 按优先级尝试各类映射
        en = self.translate_appearance(text_cn)
        if en:
            return en
        en = self.translate_ingredient(text_cn)
        if en:
            return en
        en = self.translate_ion_type(text_cn)
        if en:
            return en
        en = self.translate_product_name(text_cn)
        if en:
            return en
        # 找不到时，使用 argostranslate 翻译
        return self._translate_with_argos(text_cn)

    def _update_english_product_name(self, doc: Document, new_name: str):
        """更新产品名称为英文（保留原始段落格式和字体）"""
        for para in doc.paragraphs:
            text = para.text
            if "产品名称：" in text:
                idx = text.find("产品名称：")
                prefix = text[:idx + len("产品名称：")]
                # 保存原始段落格式
                pf = para.paragraph_format
                saved_line_spacing = pf.line_spacing
                saved_first_indent = pf.first_line_indent
                saved_left_indent = pf.left_indent
                first_run = para.runs[0] if para.runs else None
                orig_font_name = first_run.font.name if first_run else None
                orig_font_size = first_run.font.size if first_run else None
                # 清空并重建
                for run in para.runs:
                    run.text = ""
                para.clear()
                new_run = para.add_run(f"{prefix}{new_name}")
                if orig_font_name:
                    new_run.font.name = orig_font_name
                elif orig_font_size:
                    new_run.font.size = orig_font_size
                pf.line_spacing = saved_line_spacing
                pf.first_line_indent = saved_first_indent
                pf.left_indent = saved_left_indent
                return

    def _update_english_composition_table(self, doc: Document, composition_en: list):
        """更新成分表格为英文（保留原有表格样式）"""
        for table in doc.tables:
            if not table.rows:
                continue
            first_row_text = "".join(c.text for c in table.rows[0].cells)
            # 检查是否是中文成分表
            if "组分" in first_row_text or "成分" in first_row_text:
                # 先把表头改成英文
                hdr_cells = table.rows[0].cells
                if len(hdr_cells) >= 3:
                    headers_en = ["Component", "CAS NO.", "W/W %"]
                    for i, cell in enumerate(hdr_cells[:3]):
                        if cell.paragraphs and cell.paragraphs[0].runs:
                            for run in cell.paragraphs[0].runs:
                                run.text = headers_en[i] if i < len(headers_en) else ""
                        else:
                            cell.text = headers_en[i] if i < len(headers_en) else ""

                # 获取原始单元格字体样式
                sample_row = table.rows[1] if len(table.rows) > 1 else None
                sample_fonts = []
                if sample_row:
                    for cell in sample_row.cells:
                        if cell.paragraphs and cell.paragraphs[0].runs:
                            first_run = cell.paragraphs[0].runs[0]
                            sample_fonts.append({
                                "name": first_run.font.name,
                                "size": first_run.font.size,
                            })
                        else:
                            sample_fonts.append({"name": None, "size": None})

                # 更新每一行
                for i, item in enumerate(composition_en):
                    row_idx = i + 1
                    if row_idx >= len(table.rows):
                        # 现有行不够，新增一行
                        new_row = table.add_row()
                        if table.rows and len(table.rows) > 1:
                            last_row = table.rows[-2]
                            for col_idx in range(min(len(new_row.cells), len(last_row.cells))):
                                src_cell = last_row.cells[col_idx]
                                dst_cell = new_row.cells[col_idx]
                                if src_cell.paragraphs and dst_cell.paragraphs:
                                    src_para = src_cell.paragraphs[0]
                                    dst_para = dst_cell.paragraphs[0] if dst_cell.paragraphs else dst_cell.add_paragraph()
                                    src_fmt = src_para.paragraph_format
                                    dst_fmt = dst_para.paragraph_format
                                    dst_fmt.alignment = src_fmt.alignment
                                    dst_fmt.line_spacing = src_fmt.line_spacing
                                    dst_fmt.line_spacing_rule = src_fmt.line_spacing_rule
                                    dst_fmt.space_before = src_fmt.space_before
                                    dst_fmt.space_after = src_fmt.space_after
                                    dst_fmt.left_indent = src_fmt.left_indent
                                    dst_fmt.right_indent = src_fmt.right_indent
                                    dst_fmt.first_line_indent = src_fmt.first_line_indent
                                    if src_para.runs:
                                        src_run = src_para.runs[0]
                                        dst_run = dst_para.add_run("")
                                        dst_run.font.name = src_run.font.name
                                        dst_run.font.size = src_run.font.size
                                        dst_run.font.bold = src_run.font.bold
                        row = new_row
                    else:
                        row = table.rows[row_idx]

                    if len(row.cells) >= 3:
                        new_values = [
                            item.get("component_en", ""),
                            item.get("cas", ""),
                            item.get("percentage", ""),
                        ]
                        for col_idx, new_val in enumerate(new_values):
                            if col_idx >= len(row.cells):
                                continue
                            cell = row.cells[col_idx]
                            if cell.paragraphs and cell.paragraphs[0].runs:
                                para = cell.paragraphs[0]
                                for r_idx, run in enumerate(para.runs):
                                    if r_idx == 0:
                                        run.text = new_val
                                    else:
                                        run.text = ""
                            else:
                                para = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                                new_run = para.add_run(new_val)
                                if col_idx < len(sample_fonts) and sample_fonts[col_idx]["name"]:
                                    new_run.font.name = sample_fonts[col_idx]["name"]
                                if col_idx < len(sample_fonts) and sample_fonts[col_idx]["size"]:
                                    new_run.font.size = sample_fonts[col_idx]["size"]
                return

    def _update_english_physicochemical(self, doc: Document, physicochemical: dict):
        """更新理化特性为英文（保留原始段落格式和字体）"""
        # 中文 key -> 英文 key 映射
        key_map_cn_to_en = {
            "外观与性状": "Appearance",
            "离子性": "Ionicity",
            "离子型": "Ionicity",
            "PH值": "pH",
            "熔点": "Melting point",
            "沸点": "Boiling point",
            "沸点/沸点范围（℃）": "Boiling point",
            "相对密度": "Density",
            "闪点": "Flash Point",
            "溶解性": "Solubility",
        }

        for para in doc.paragraphs:
            text = para.text
            for cn_key, en_key in key_map_cn_to_en.items():
                if cn_key not in text:
                    continue
                # 找到冒号位置
                idx = text.find(cn_key)
                if idx < 0:
                    continue
                colon_idx = text.find("：", idx)
                if colon_idx < 0:
                    continue
                value_cn = text[colon_idx + 1:].strip()
                if not value_cn:
                    continue

                # 翻译值
                value_en = value_cn
                if cn_key in ["外观与性状"]:
                    value_en = self.translate_appearance(value_cn) or self.translate_text_fallback(value_cn)
                elif cn_key in ["离子性", "离子型"]:
                    value_en = self.translate_ion_type(value_cn) or self.translate_text_fallback(value_cn)
                else:
                    # 数值类的保留原值（仅做单位转换等简单处理）
                    value_en = value_cn

                # 保存原始段落格式
                pf = para.paragraph_format
                saved_line_spacing = pf.line_spacing
                saved_first_indent = pf.first_line_indent
                saved_left_indent = pf.left_indent
                first_run = para.runs[0] if para.runs else None
                orig_font_name = first_run.font.name if first_run else None
                orig_font_size = first_run.font.size if first_run else None
                # 清空并重建
                for run in para.runs:
                    run.text = ""
                para.clear()
                new_run = para.add_run(f"{en_key}：{value_en}")
                if orig_font_name:
                    new_run.font.name = orig_font_name
                elif orig_font_size:
                    new_run.font.size = orig_font_size
                pf.line_spacing = saved_line_spacing
                pf.first_line_indent = saved_first_indent
                pf.left_indent = saved_left_indent
                break

    def _update_english_section_titles(self, doc: Document):
        """将文档中的中文章节标题翻译为英文"""
        import re

        # 主标题翻译
        MAIN_TITLE_CN = "化学品安全资料说明书"
        MAIN_TITLE_EN = "MATERIAL SAFETY DATA SHEET"

        # 中文章节标题 -> 英文章节标题映射
        section_map = {
            "化学品安全资料说明书": MAIN_TITLE_EN,
            "化学品及标识": "1. Chemicals and Identification",
            "化学品及其它标识": "1. Chemicals and Identification",
            "危险性概述": "2. Hazards Identification",
            "成分/组成资讯": "3. Composition Information",
            "成分/组成资料": "3. Composition Information",
            "急救措施": "4. First Aid Measures",
            "消防措施": "5. Fire Protection Measures",
            "泄漏应急处理": "6. Accidental Release Measures",
            "事故泄漏应急处理": "6. Accidental Release Measures",
            "操作处置与储存": "7. Handling and Storage",
            "接触控制/个人防护": "8. Exposure Controls/Personal Protection",
            "理化特性": "9. Physical and Chemical Properties",
            "稳定性和反应活性": "10. Stability and Reactivity",
            "毒性学资料": "11. Toxicological Information",
            "生态学资料": "12. Ecological Information",
            "生态资料": "12. Ecological Information",
            "废弃处理": "13. Disposal Considerations",
            "运输资讯": "14. Transport Information",
            "运输信息": "14. Transport Information",
            "法规资讯": "15. Regulatory Information",
            "法规信息": "15. Regulatory Information",
            "其它资讯": "16. Additional Information",
            "其他资讯": "16. Additional Information",
        }

        # 统计主标题出现次数，用于判断是否是"中文主标题 + 英文主标题"的组合
        main_title_count = 0
        for para in doc.paragraphs:
            if MAIN_TITLE_CN in para.text or MAIN_TITLE_EN in para.text:
                main_title_count += 1

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # --- 主标题处理 ---
            # 场景A：段落同时含中英文（只有中文主标题），替换为纯英文
            # 场景B：有两个独立段落分别含中/英文主标题（原文就是中英分离的），
            #        只处理含MAIN_TITLE_CN的段落，将其替换为纯英文；英文段落保持不变
            if MAIN_TITLE_CN in text and MAIN_TITLE_EN not in text:
                # 只有中文主标题 → 替换为纯英文
                self._replace_para_text(para, MAIN_TITLE_EN)
                continue
            if MAIN_TITLE_CN not in text and MAIN_TITLE_EN in text:
                # 只有英文主标题 → 保持原样（如果有两个独立段落，这是英文那个）
                if main_title_count >= 2:
                    # 有中英文分离的两段，英文段落保持不变
                    continue
                else:
                    # 只有一个英文段落，说明原文已经是纯英文了
                    continue

            # --- 章节标题处理 ---
            # 去掉数字前缀：全角数字(１-９)、半角数字(1-9)、中文数字(一-九)
            # + 可选的 '.' 或 '、' 或 '、' + 空格
            text_stripped = re.sub(
                r'^[0-9１-９一-鿿]+[.、\s]+',
                '',
                text
            )
            # 去除尾随的全角冒号 "：" 和半角冒号 ":" 以及空格
            text_stripped = text_stripped.rstrip('：: ​')

            replaced = False
            for cn_title, en_title in section_map.items():
                if cn_title == MAIN_TITLE_CN:
                    continue  # 主标题已在上面处理

                # 用去掉前缀+尾随冒号后的文本匹配
                if text_stripped == cn_title:
                    # 替换整段：去掉旧标题，换成英文标题
                    # 保留段落中标题之后的内容
                    after_title = text[len(cn_title):]
                    self._replace_para_text(para, en_title + after_title)
                    replaced = True
                    break

                # 也检查原文（去除前缀后英文残留的情况，如 "15、15. Regulatory Information："
                # 去掉前缀后 text_stripped 为 "15. Regulatory Information："，
                # 与 key "法规资讯" 不匹配，但与英文标题直接重叠）
                if en_title in text and text_stripped.startswith(en_title.rstrip('：')):
                    # 英文标题已存在但中文前缀没去掉整段 → 去除中文章节标题
                    # 如 "15、15. Regulatory Information：" → "15. Regulatory Information："
                    after = text[len(cn_title):]
                    self._replace_para_text(para, en_title + after)
                    replaced = True
                    break

                # 直接原文匹配
                if text == cn_title or text.startswith(cn_title + '：') or text.startswith(cn_title + ' '):
                    after = text[len(cn_title):]
                    self._replace_para_text(para, en_title + after)
                    replaced = True
                    break

    def _update_english_header(self, doc: Document, msds_number: str, revision_date: str):
        """更新页眉（MSDS编号和修订时间保持原样，仅确保格式正确）"""
        for section in doc.sections:
            for para in section.header.paragraphs:
                text = para.text
                if not text:
                    continue

                pf = para.paragraph_format
                saved_line_spacing = pf.line_spacing
                saved_first_indent = pf.first_line_indent
                saved_left_indent = pf.left_indent
                first_run = para.runs[0] if para.runs else None
                orig_font_name = first_run.font.name if first_run else None
                orig_font_size = first_run.font.size if first_run else None

                new_text = text

                # 替换 MSDS 编号
                if msds_number and "MSDS编号：" in text:
                    for marker in ["修定时间", "修订时间", "修订日期"]:
                        pos = text.find(marker)
                        if pos >= 0:
                            msds_prefix = text.split("MSDS编号：")[0] + "MSDS编号："
                            m = re.search(r'(\d{4}/\d{2}/\d{2})', text)
                            orig_date = m.group(1) if m else ""
                            date_to_use = revision_date if revision_date else orig_date
                            new_text = f"{msds_prefix}{msds_number}   {marker}：{date_to_use}"
                            break

                # 替换修订时间
                if revision_date:
                    m = re.search(r'(\d{4}/\d{2}/\d{2})', new_text)
                    if m:
                        new_text = re.sub(r'(\d{4}/\d{2}/\d{2})', revision_date, new_text, count=1)

                if new_text != text:
                    # 检测段落是否包含绘图（如 logo/印章）
                    from lxml import etree
                    para_xml = para._element.xml
                    has_drawing = '<w:drawing' in para_xml

                    if has_drawing:
                        # 含绘图的段落：只清空文字run，保留绘图run
                        wns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                        for r in para._element.findall(f'{{{wns}}}r'):
                            has_r_drawing = False
                            if r.find(f'{{{wns}}}drawing') is not None:
                                has_r_drawing = True
                            else:
                                r_xml = etree.tostring(r, encoding='unicode')
                                if '<w:drawing' in r_xml or '<w:anchor' in r_xml:
                                    has_r_drawing = True
                            if not has_r_drawing:
                                for t in r.findall(f'{{{wns}}}t'):
                                    t.text = ''
                        # 在末尾添加新文字run
                        new_run = para.add_run(new_text)
                        if orig_font_name:
                            new_run.font.name = orig_font_name
                        elif orig_font_size:
                            new_run.font.size = orig_font_size
                    else:
                        # 普通段落：完全替换
                        for run in para.runs:
                            run.text = ""
                        para.clear()
                        new_run = para.add_run(new_text)
                        if orig_font_name:
                            new_run.font.name = orig_font_name
                        elif orig_font_size:
                            new_run.font.size = orig_font_size

                    pf.line_spacing = saved_line_spacing
                    pf.first_line_indent = saved_first_indent
                    pf.left_indent = saved_left_indent

    def _update_english_section16(self, doc: Document, revision_date: str):
        """更新第16节"其他资讯"中的版次和更新日期"""
        if not revision_date:
            return

        parts = revision_date.split("/")
        if len(parts) < 2:
            return
        year = parts[0]
        month = parts[1]
        banci = f"{year}-{month}"
        gengxin = revision_date

        for para in doc.paragraphs:
            text = para.text
            new_text = text

            # 先翻译标签：版次 -> Revision No.，更新日期 -> Revision Date
            new_text = new_text.replace("版次", "Revision No.")
            new_text = new_text.replace("更新日期", "Revision Date")
            new_text = new_text.replace("修改说明", "Revision Note")
            new_text = new_text.replace("填表部门", "Prepared By")
            new_text = new_text.replace("其它说明", "Other Information")

            # 替换版次日期
            m_ban = re.search(r'(Revision No[.:]\s*)([\d]{4}-[\d]{2})', new_text)
            if m_ban:
                new_text = new_text[:m_ban.start()] + m_ban.group(1) + banci + new_text[m_ban.end():]

            # 替换更新日期
            m_geng = re.search(r'(Revision Date[.:]\s*)([\d]{4}/[\d]{1,2}/[\d]{1,2})', new_text)
            if m_geng:
                new_text = new_text[:m_geng.start()] + m_geng.group(1) + gengxin + new_text[m_geng.end():]

            if new_text != text:
                self._replace_para_text(para, new_text)

    def _update_section16_values(self, doc: Document, revision_date: str):
        """
        在通用翻译完成后，更新 Section 16 的版次和更新日期数值。
        标签已由 _translate_paragraph_text() 翻译为英文（Revision No.、Revision Date 等），
        此函数只替换对应的日期数值。
        """
        if not revision_date:
            return

        parts = revision_date.split("/")
        if len(parts) < 2:
            return
        year = parts[0]
        month = parts[1]
        banci = f"{year}-{month}"      # 版次数值: 2025-12

        for para in doc.paragraphs:
            text = para.text
            new_text = text

            # 替换版次数值（Revision: 后面跟着的 YYYY-MM 或 JAN.,YYYY 等）
            m_ban = re.search(r'(Revision:\s*)([\d]{4}-[\d]{2}|[A-Za-z.,\s]+)', new_text)
            if m_ban:
                new_text = new_text[:m_ban.start()] + m_ban.group(1) + banci + new_text[m_ban.end():]

            # 替换更新日期数值（Update Date: 后面的所有内容，包含中文日期格式）
            # 支持：Update Date:2026年03月 或 Update Date:JAN,2026 等
            m_geng = re.search(r'(Update Date:\s*)(.+)', new_text)
            if m_geng:
                new_text = new_text[:m_geng.start()] + m_geng.group(1) + revision_date + new_text[m_geng.end():]

            if new_text != text:
                self._replace_para_text(para, new_text)

    def _translate_physicochemical_values(self, doc: Document):
        """
        翻译理化特性字段的值（Appearance、Ionicity、Solubility 等）。
        在通用段落翻译之后调用，此时字段名已是英文。
        使用 appearance_color_mapping 和 ion_type_mapping 做值翻译。
        """
        # 字段名（英文）→ 翻译函数
        value_translators = {
            "Appearance": self.translate_appearance,
            "Ionicity": self.translate_ion_type,
            "Solubility": self.translate_appearance,
        }

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            for en_field, translator in value_translators.items():
                if en_field + ": " in text or en_field + ":" in text:
                    # 提取冒号后面的值
                    parts = text.split(": ", 1)
                    if len(parts) < 2:
                        parts = text.split(":", 1)
                    if len(parts) < 2:
                        continue
                    field_part = parts[0].strip()
                    value_part = parts[1].strip()
                    if field_part != en_field:
                        continue
                    # 翻译值
                    if value_part and value_part not in ("No data", "Not applicable", "None"):
                        translated_value = translator(value_part)
                        if translated_value and translated_value != value_part:
                            new_text = f"{en_field}: {translated_value}"
                            self._replace_para_text(para, new_text)
                    break

    def _translate_product_name_field(self, doc: Document, product_name_en: str):
        """
        翻译产品名称字段值。
        在通用段落翻译之后调用，此时字段名已是英文 "Product Name"。
        将 "Product Name: 有机硅柔软剂" 替换为 "Product Name: SILICONE SOFTENER"。
        """
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            # 匹配 "Product Name: 中文名" 或 "Product Name：中文名"
            for sep in [": ", "："]:
                if f"Product Name{sep}" in text:
                    idx = text.index(f"Product Name{sep}") + len(f"Product Name{sep}")
                    # 如果当前值已经是英文（或接近英文），跳过
                    current_value = text[idx:].strip()
                    if current_value and current_value == product_name_en:
                        return
                    # 检查当前值是否需要翻译
                    if current_value != product_name_en:
                        new_text = f"Product Name: {product_name_en}"
                        self._replace_para_text(para, new_text)
                    return

    def _update_physicochemical(self, doc: Document, physicochemical: dict):
        """
        更新理化特性，保留原始段落格式和字体。
        格式：{"外观与性状": "蓝色透明液体", "离子性": "阴离子", "PH值": "3±1", ...}
        """
        from docx.shared import Pt

        # 映射前端字段名到文档中实际存在的 key
        key_map = {
            'appearance': '外观与性状',
            'ion_type': '离子性',
            'ph': 'PH值',
            'melting_point': '熔点',
            'boiling_point': '沸点',
            'density': '相对密度',
            'flash_point': '闪点',
            'solubility': '溶解性',
        }

        for para in doc.paragraphs:
            text = para.text
            for field_key, value in physicochemical.items():
                if not value:
                    continue
                # 找到文档中对应的 key
                doc_key = key_map.get(field_key, field_key)
                # 大小写敏感匹配
                if f"{doc_key}：" not in text:
                    continue

                # 找到冒号位置（用 doc_key 本身在文本中的位置）
                colon_idx = text.find(doc_key)
                if colon_idx == -1:
                    continue
                colon_idx = text.find("：", colon_idx)
                if colon_idx == -1:
                    continue
                prefix = text[:colon_idx + 1]  # 保留原始 key 大小写
                new_text = f"{prefix}{value}"

                # 保存原始段落格式
                pf = para.paragraph_format
                saved_line_spacing = pf.line_spacing
                saved_first_indent = pf.first_line_indent
                saved_left_indent = pf.left_indent

                # 保存原始 run 的字体属性
                first_run = para.runs[0] if para.runs else None
                orig_font_name = first_run.font.name if first_run else None
                orig_font_size = first_run.font.size if first_run else None

                # 清空所有 run 并重建
                for run in para.runs:
                    run.text = ""
                para.clear()

                new_run = para.add_run(new_text)
                # 保持与原始一致的字体
                if orig_font_name:
                    new_run.font.name = orig_font_name
                elif orig_font_size:
                    new_run.font.size = orig_font_size

                # 恢复段落格式
                pf.line_spacing = saved_line_spacing
                pf.first_line_indent = saved_first_indent
                pf.left_indent = saved_left_indent
                break


    # ===================== Template-based MSDS Generation =====================

    def generate_msds_from_template(
        self,
        ledger_data: dict,
        language: str,
        msds_number: str,
        revision_date: str
    ) -> tuple:
        """
        Generate MSDS from unified template (cnMSDS.docx or enMSDS.docx).
        
        Args:
            ledger_data: dict with keys: customs_name, appearance, ion_type, ph, composition,
                         product_name_en, appearance_en, ion_type_en
            language: "cn" or "en"
            msds_number: MSDS编号
            revision_date: 修订时间 (YYYY/MM/DD)
        
        Returns:
            (docx_bytes, doc_key)
        """
        import time
        from docx import Document
        
        # 1. Select template
        template_key = "msds" if language == "cn" else "msds_en"
        template_path = TEMPLATES.get(template_key)
        if not template_path or not os.path.exists(template_path):
            raise FileNotFoundError(f"MSDS template not found: {template_key}")
        
        doc = Document(template_path)
        
        # 2. Fill product info
        if language == "cn":
            self._fill_placeholder(doc, "{{customs_name}}", ledger_data.get("customs_name", ""))
            self._fill_placeholder(doc, "{{appearance}}", ledger_data.get("appearance", "无资料"))
            self._fill_placeholder(doc, "{{ion_type}}", ledger_data.get("ion_type", "无资料"))
            self._fill_placeholder(doc, "{{ph}}", ledger_data.get("ph", "无资料"))
        else:
            # Auto-fill English fields from mapping files if empty
            ledger_data = self._auto_fill_english_fields(ledger_data)
            self._fill_placeholder(doc, "{{product_name_en}}", ledger_data.get("product_name_en", ""))
            self._fill_placeholder(doc, "{{appearance_en}}", ledger_data.get("appearance_en", "No data"))
            self._fill_placeholder(doc, "{{ion_type_en}}", ledger_data.get("ion_type_en", "No data"))
            self._fill_placeholder(doc, "{{ph}}", ledger_data.get("ph", "No data"))
        
        # 3. Fill composition table
        composition = ledger_data.get("composition", [])
        if composition:
            self._fill_composition_from_template(doc, composition, language)
        
        # 4. Fill header (MSDS编号 + 修订时间)
        self._fill_placeholder(doc, "{{msds_number}}", msds_number)
        self._fill_placeholder(doc, "{{revision_date}}", revision_date)
        
        # 5. Fill Section 16 (版次 + 更新日期)
        parts = revision_date.split("/") if revision_date else ["", ""]
        if len(parts) >= 2:
            year = parts[0]
            month = parts[1]
            revision = f"{year}-{month}"
            update_date_cn = f"{year}年{int(month)}月"
            update_date_en = f"{year}/{month}"
        else:
            revision = ""
            update_date_cn = ""
            update_date_en = ""
        
        self._fill_placeholder(doc, "{{revision}}", revision)
        if language == "cn":
            self._fill_placeholder(doc, "{{update_date}}", update_date_cn)
        else:
            self._fill_placeholder(doc, "{{update_date}}", update_date_en)
        
        # 6. Save
        from io import BytesIO
        buf = BytesIO()
        doc.save(buf)
        content = buf.getvalue()
        
        product_name = ledger_data.get("customs_name", ledger_data.get("product_name_en", "unknown"))
        doc_key = f"msds_{product_name}_{int(time.time())}"
        
        return content, doc_key

    def _auto_fill_english_fields(self, ledger_data: dict) -> dict:
        """
        Auto-fill English fields from mapping files if empty.
        Preserves existing values — only fills blanks.
        """
        # Product name: products_name_mapping.json
        if not ledger_data.get("product_name_en"):
            cn_name = ledger_data.get("customs_name", "")
            en_name = self.translate_product_name(cn_name)
            if en_name:
                ledger_data["product_name_en"] = en_name

        # Appearance: appearance_color_mapping.json
        if not ledger_data.get("appearance_en"):
            cn_appearance = ledger_data.get("appearance", "")
            en_appearance = self._appearance_map.get(cn_appearance)
            if en_appearance:
                ledger_data["appearance_en"] = en_appearance

        # Ion type: ion_type_mapping.json
        if not ledger_data.get("ion_type_en"):
            cn_ion = ledger_data.get("ion_type", "")
            en_ion = self._ion_type_map.get(cn_ion)
            if en_ion:
                ledger_data["ion_type_en"] = en_ion

        # Composition: translate component_cn to component_en
        composition = ledger_data.get("composition", [])
        for item in composition:
            if not item.get("component_en") and item.get("component_cn"):
                item["component_en"] = self.translate_text_fallback(item["component_cn"])

        return ledger_data

    @staticmethod
    def convert_docx_to_pdf(docx_bytes: bytes) -> tuple:
        """
        Convert docx bytes to PDF using Word/WPS COM on Windows.
        Falls back to returning original docx if conversion fails.
        
        Returns: (content_bytes, extension)
        """
        import win32com.client
        import tempfile
        
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            
            with tempfile.TemporaryDirectory() as tmpdir:
                docx_path = os.path.join(tmpdir, "input.docx")
                pdf_path = os.path.join(tmpdir, "output.pdf")
                
                with open(docx_path, "wb") as f:
                    f.write(docx_bytes)
                
                doc = word.Documents.Open(os.path.abspath(docx_path))
                doc.SaveAs2(os.path.abspath(pdf_path), FileFormat=17)  # 17 = wdFormatPDF
                doc.Close(0)
                
                # Read PDF before attempting to quit Word
                if os.path.exists(pdf_path) and os.path.getsize(pdf_path) > 0:
                    with open(pdf_path, "rb") as f:
                        pdf_bytes = f.read()
                    return pdf_bytes, ".pdf"
            
            return docx_bytes, ".docx"
        except Exception:
            return docx_bytes, ".docx"
        finally:
            try:
                word.Quit()
            except Exception:
                pass

    def _fill_placeholder(self, doc, placeholder, value):
        """Replace placeholder text in all paragraphs, preserving surrounding text."""
        # Collect all paragraphs to scan: body + headers + footers
        all_paras = list(doc.paragraphs)
        for section in doc.sections:
            all_paras.extend(section.header.paragraphs)
            all_paras.extend(section.footer.paragraphs)

        for para in all_paras:
            if placeholder in para.text:
                full_text = para.text.replace(placeholder, value)
                if para.runs:
                    first_run = para.runs[0]
                    orig_font_name = first_run.font.name
                    orig_font_size = first_run.font.size
                    for run in para.runs:
                        run.text = ''
                    para.runs[0].text = full_text
                    if orig_font_name:
                        para.runs[0].font.name = orig_font_name
                    if orig_font_size:
                        para.runs[0].font.size = orig_font_size

        # Also check in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if placeholder in para.text:
                            if para.runs:
                                full_text = para.text.replace(placeholder, value)
                                for run in para.runs:
                                    run.text = ''
                                para.runs[0].text = full_text

    def _fill_composition_from_template(self, doc, composition, language):
        """Fill composition table from template placeholder."""
        if not doc.tables:
            return
        
        table = doc.tables[0]
        
        # Remove placeholder row(s) except header
        while len(table.rows) > 1:
            row = table.rows[-1]
            row._tr.getparent().remove(row._tr)
        
        # Add rows for each component
        for item in composition:
            new_row = table.add_row()
            if language == "cn":
                new_row.cells[0].text = item.get("component_cn", "")
            else:
                new_row.cells[0].text = item.get("component_en", "") or item.get("component_cn", "")
            new_row.cells[1].text = item.get("cas", "")
            new_row.cells[2].text = item.get("percentage", "")

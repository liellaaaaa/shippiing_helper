# backend/app/services/document_service.py
import openpyxl
import xlrd
from typing import Tuple, Optional
import os, time, base64, uuid
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.core.config import TEMPLATES, MSDS_DIR
from app.services.msds_service import MSDSService
from app.database import SessionLocal
from app.models.order import Order, OrderItem
from app.models.order_pi_record import OrderPiRecord
from app.models.pi_contract import PiContract, PiContractItem
from app.services.customs_declaration_service import CustomsDeclarationService


import re

# XML 1.0 compatible character set: U+0009, U+000A, U+000D, U+0020–U+10FFFF
_XML_ILLEGALChars = re.compile(
    r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x84\x86-\x9f\ud800-\udfff￾￿]"
)


def _sanitize_for_xml(text: str) -> str:
    """Remove XML-illegal control characters from text."""
    return _XML_ILLEGALChars.sub("", text)


def _parse_date(value: str):
    """将多种格式的日期字符串转为 Python datetime，失败时返回原字符串。"""
    if not value:
        return value
    for fmt in ("%Y-%m-%d", "%Y-%m-%d %H:%M:%S", "%Y年%m月%d日"):
        try:
            return datetime.strptime(value[:len("2026-12-31")], fmt)
        except Exception:
            pass
    return value


def convert_doc_to_docx(doc_path: str) -> bytes:
    """
    将 .doc 文件（binary，非 OLE）转换为 .docx 格式。
    使用 antiword 提取纯文本，再按段落写入 python-docx 文档。
    返回 docx 内容的 bytes。
    """
    text = _extract_doc_text(doc_path)
    doc = Document()
    for line in text.split("\n"):
        line = _sanitize_for_xml(line.rstrip())
        if not line:
            doc.add_paragraph("")
        else:
            doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _extract_doc_text(doc_path: str) -> str:
    """
    从 .doc 文件提取纯文本（antiword 优先，chardet 回退）。
    尝试多种编码以确保中文正确提取。
    """
    import subprocess, chardet

    # 优先尝试 antiword（可能返回 utf-8 或 latin-1 编码）
    for encoding in ("utf-8", "latin-1", "cp1252", "gb2312", "gbk"):
        try:
            result = subprocess.run(
                ["antiword", "-f", doc_path] if encoding != "utf-8" else ["antiword", doc_path],
                capture_output=True,
                timeout=30,
            )
            if result.returncode == 0 and result.stdout:
                try:
                    text = result.stdout.decode(encoding, errors="strict")
                except UnicodeDecodeError:
                    text = result.stdout.decode(encoding, errors="replace")
                # 检查是否提取到了有效内容（包含非ASCII字符或足够多的ASCII）
                if text and len(text.strip()) > 10:
                    return text
        except (subprocess.SubprocessError, FileNotFoundError):
            break
        except Exception:
            continue

    # 回退：读取二进制文件，用 chardet 检测编码
    try:
        with open(doc_path, "rb") as f:
            raw = f.read()
        detected = chardet.detect(raw)
        encoding = detected.get("encoding", "latin-1") or "latin-1"
        return raw.decode(encoding, errors="replace")
    except Exception:
        return ""


def convert_xls_to_xlsx(template_path: str) -> bytes:
    """用 xlrd 读取 .xls，用 openpyxl 写出 .xlsx，保持单元格数据不变。"""
    wb_xls = xlrd.open_workbook(template_path)
    ws_xls = wb_xls.sheet_by_index(0)
    wb = openpyxl.Workbook()
    ws = wb.active
    for r in range(ws_xls.nrows):
        for c in range(ws_xls.ncols):
            ws.cell(r + 1, c + 1).value = ws_xls.cell_value(r, c)
    buf = BytesIO()
    wb.save(buf)
    return buf.getvalue()


def find_marker_cell(ws, marker_text: str) -> Tuple[int, int]:
    """在 worksheet 中搜索 {{MARK_XXX}} 锚点，返回 (row, column)，1-based。"""
    for row in ws.iter_rows():
        for cell in row:
            if cell.value and str(cell.value).strip() == marker_text:
                return cell.row, cell.column
    raise ValueError(f"Marker '{marker_text}' not found in template")


def clear_markers(ws):
    """渲染后清除所有 {{MARK_XXX}} 锚点。"""
    for row in ws.iter_rows():
        for cell in row:
            if isinstance(cell.value, str) and cell.value.startswith("{{MARK_"):
                cell.value = None


def replace_placeholder(ws, placeholder: str, value) -> bool:
    """在 worksheet 中查找 {{占位符}} 并替换为 value。返回是否替换成功。"""
    if value is None:
        return False
    for row in ws.iter_rows():
        for cell in row:
            if cell.value and str(cell.value).strip() == placeholder:
                cell.value = value
                return True
    return False


def _to_invoice_no(contract_no: str) -> str:
    """合同号转发票号：去掉字母前缀（HT/HH/HTPK等），换成 IN。
    例: HT260304E01→IN260304E01, HTPK260304→IN260304, HH12345→IN12345
    """
    if not contract_no:
        return ""
    # 找到第一个数字的位置，取后面的部分
    m = re.search(r'\d', contract_no)
    if m:
        return "IN" + contract_no[m.start():]
    # 没有数字，直接加 IN 前缀
    return "IN" + contract_no


def _amount_to_chinese_upper(amount: float) -> str:
    """将金额转换为中文大写（元整）。
    例: 49720 → '肆万玖仟柒佰贰拾元整'
    """
    digits = "零壹贰叁肆伍陆柒捌玖"
    units = ["", "拾", "佰", "仟"]
    big_units = ["", "万", "亿"]

    if amount == 0:
        return "零元整"

    amount = round(amount, 2)
    int_part = int(amount)
    dec_part = round((amount - int_part) * 100)

    if int_part > 0:
        int_str = str(int_part)
        length = len(int_str)
        result = ""
        prev_zero = False

        for i, ch in enumerate(int_str):
            digit = int(ch)
            pos = length - 1 - i
            big_unit_idx = pos // 4
            unit_idx = pos % 4

            if digit != 0:
                if prev_zero:
                    result += "零"
                result += digits[digit] + units[unit_idx]
                prev_zero = False
            else:
                prev_zero = True

            if unit_idx == 0 and big_unit_idx > 0:
                result = result.rstrip("零")
                result += big_units[big_unit_idx]
                prev_zero = False

        result = result.rstrip("零")
        result += "元"
    else:
        result = "零元"

    if dec_part > 0:
        jiao = dec_part // 10
        fen = dec_part % 10
        if jiao > 0:
            result += digits[jiao] + "角"
        if fen > 0:
            result += digits[fen] + "分"
    else:
        result += "整"

    return result


class DocumentService:
    def __init__(self):
        self.msds_service = MSDSService()

    def fill_booking_template(self, fields: dict, template_key: str = "booking_marked") -> bytes:
        """
        直接操作 xlsx zip 内部 XML，精准替换 {{FIELD_NAME}} 单元格值，
        不经过 openpyxl 加载/保存，完整保留所有 shapes 和格式。
        fields: dict，键为 FIELD_NAME（不带花括号），值为字符串
        template_key: config.py 中的 TEMPLATES 键，默认 "booking_marked"
        """
        import zipfile, re, shutil, os
        from io import BytesIO
        from lxml import etree

        template_path = TEMPLATES.get(template_key, TEMPLATES.get("booking_marked", TEMPLATES.get("booking")))
        NS = 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'

        # 1. 先用 openpyxl 只读模式找出 {{FIELD}} 单元格坐标（不保存）
        wb = openpyxl.load_workbook(template_path, data_only=True)
        ws = wb.active
        sheet_name = ws.title
        marker_map = {}  # cell_ref -> field_key
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    m = re.match(r"^\{\{(\w+)\}\}$", str(cell.value).strip())
                    if m:
                        marker_map[cell.coordinate] = m.group(1)
        wb.close()

        # 2. 直接修改 xlsx zip 内的 sheet XML
        tmp_path = template_path + ".tmp"
        shutil.copy(template_path, tmp_path)

        with zipfile.ZipFile(tmp_path, 'r') as zin:
            # 找到工作表文件路径
            wb_xml = etree.fromstring(zin.read('xl/workbook.xml'))
            sheets_el = wb_xml.find(f'{{{NS}}}sheets')
            r_id = None
            for s in sheets_el:
                if s.get('name') == sheet_name:
                    r_id = s.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}id')
                    break

            rels_xml = etree.fromstring(zin.read('xl/_rels/workbook.xml.rels'))
            sheet_file = None
            for rel in rels_xml:
                if rel.get('Id') == r_id:
                    sheet_file = 'xl/' + rel.get('Target')
                    break

            sheet_tree = etree.fromstring(zin.read(sheet_file))

            # 3. 遍历单元格，替换 {{FIELD}} 的值
            for c_el in sheet_tree.iter(f'{{{NS}}}c'):
                r = c_el.get('r')
                if r in marker_map:
                    field_key = marker_map[r]
                    value = fields.get(field_key, "")
                    # 单位后缀：如果值已包含单位则不追加
                    if field_key == "NO_KIND_PKG" and value:
                        val_upper = str(value).upper()
                        if not any(unit in val_upper for unit in ["PALLETS", "DRUMS", "DRUM", "CTNS", "PCS", "BUNDLES"]):
                            value = f"{value} PALLETS"
                    elif field_key == "MEASUREMENT" and value:
                        if "CBM" not in str(value).upper():
                            value = f"{value} CBM"
                    # 写回 <v> 元素，并改为内联字符串类型（t="str"）
                    # 原始单元格可能是 t="s"（共享字符串索引），直接改值会导致查表错误
                    v_el = c_el.find(f'{{{NS}}}v')
                    if v_el is None:
                        v_el = etree.SubElement(c_el, f'{{{NS}}}v')
                    v_el.text = str(value)
                    # 改为内联字符串，避免共享字符串表索引错误
                    c_el.set('t', 'str')

            # 4. 重新打包 zip
            out = BytesIO()
            with zipfile.ZipFile(out, 'w', zipfile.ZIP_DEFLATED) as zout:
                for item in zin.namelist():
                    if item == sheet_file:
                        zout.writestr(item, etree.tostring(
                            sheet_tree, xml_declaration=True, encoding='UTF-8', standalone=True))
                    else:
                        zout.writestr(item, zin.read(item))

        os.remove(tmp_path)
        out.seek(0)
        return out.read()

    def generate_booking(self, fields: dict | None = None, template_type: str = "xlsx") -> Tuple[bytes, str, str]:
        """
        生成订舱单：template_type="multi" 使用多产品模板 booking_multi，
        其他情况使用 booking_marked。
        fields: 可选，键为字段名（无花括号），值未提供则返回空白模板。
        """
        if template_type == "multi":
            template_key = "booking_multi"
        else:
            template_key = "booking_marked"

        if fields:
            content_xlsx = self.fill_booking_template(fields, template_key=template_key)
        else:
            # 无参调用也使用 xlsx 已标记模板（空白版本）
            template_path = TEMPLATES.get(template_key, TEMPLATES.get("booking_marked", TEMPLATES.get("booking")))
            with open(template_path, "rb") as f:
                content_xlsx = f.read()

        doc_key = f"booking_{int(time.time())}"
        return content_xlsx, doc_key, base64.b64encode(content_xlsx).decode()

    def generate_msds(self, product_name: str) -> Tuple[bytes, str, str]:
        template_path = TEMPLATES["msds"]
        db = SessionLocal()
        try:
            msds_index = self.msds_service.get_msds_by_product(product_name, db)
            # Find a usable template: prefer standard .docx, then matched .docx, then any .docx in MSDS_DIR
            effective_template: Optional[str] = None
            if template_path.endswith(".docx") and os.path.exists(template_path):
                effective_template = template_path
            elif msds_index:
                matched_path = msds_index.file_path
                if matched_path.endswith(".docx") and os.path.exists(matched_path):
                    effective_template = matched_path
                else:
                    # .doc file is not loadable by python-docx — try to find a .docx fallback in the dir
                    msds_dir = os.path.dirname(matched_path) if os.path.exists(os.path.dirname(matched_path)) else MSDS_DIR
                    for fname in os.listdir(msds_dir):
                        if fname.endswith(".docx"):
                            candidate = os.path.join(msds_dir, fname)
                            if os.path.exists(candidate):
                                effective_template = candidate
                                break
            if not effective_template:
                # Last resort: scan MSDS_DIR for any .docx
                for fname in os.listdir(MSDS_DIR):
                    if fname.endswith(".docx"):
                        candidate = os.path.join(MSDS_DIR, fname)
                        if os.path.exists(candidate):
                            effective_template = candidate
                            break
            if not effective_template:
                # 尝试把 matched .doc 文件转成 .docx（从零构建表格结构）
                if msds_index and msds_index.file_path.endswith(".doc") and os.path.exists(msds_index.file_path):
                    content = self._build_msds_docx_from_doc(msds_index.file_path, product_name)
                    doc_key = f"msds_{product_name}_{int(time.time())}"
                    return content, doc_key, base64.b64encode(content).decode()
                else:
                    # 最后尝试：将 MSDS_DIR 中任意第一个 .doc 转换为 .docx
                    for fname in os.listdir(MSDS_DIR):
                        if fname.endswith(".doc"):
                            candidate = os.path.join(MSDS_DIR, fname)
                            if os.path.exists(candidate):
                                try:
                                    content = self._build_msds_docx_from_doc(candidate, product_name)
                                    doc_key = self._sanitize_doc_key(f"msds_{product_name}_{int(time.time())}")
                                    return content, doc_key, base64.b64encode(content).decode()
                                except Exception:
                                    continue
                    raise FileNotFoundError(f"MSDS template not found for product: {product_name}")

            doc = Document(effective_template)
            tables = doc.tables
            if msds_index and tables:
                text = self.msds_service.extract_text(msds_index.file_path)
                composition = self.msds_service.extract_composition_table(text)
                props = self.msds_service.extract_physical_props(text)
                if len(tables) >= 1:
                    t0 = tables[0]
                    t0.cell(0, 1).text = props.get("physical_form", "")
                    t0.cell(1, 1).text = props.get("ion_type", "")
                    t0.cell(2, 1).text = props.get("ph", "")
                if len(tables) >= 2:
                    comp_table = tables[1]
                    for i, item in enumerate(composition):
                        if i >= len(comp_table.rows):
                            break
                        row = comp_table.rows[i]
                        row.cells[0].text = str(i + 1)
                        row.cells[1].text = item.get("component", "")
                        row.cells[2].text = item.get("cas", "")
                        row.cells[3].text = item.get("percentage", "")
            buf = BytesIO()
            doc.save(buf)
            content = buf.getvalue()
            # Fix XML-level ordering: if a table immediately precedes a section heading
            # paragraph, swap them so the heading appears before the table in the rendered document.
            content = self._swap_table_heading_order(content)
        finally:
            db.close()
        doc_key = self._sanitize_doc_key(f"msds_{product_name}_{int(time.time())}")
        return content, doc_key, base64.b64encode(content).decode()

    @staticmethod
    def _swap_table_heading_order(content: bytes) -> bytes:
        """
        Fix OOXML body element ordering: if a <w:tbl> immediately precedes a
        <w:p> that starts with a section heading (e.g. "3、" or "13、"),
        swap them so the heading paragraph appears before the table.
        This corrects rendering in editors (OnlyOffice) that follow XML element order.
        """
        import zipfile, re
        from lxml import etree

        NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        # Section heading pattern: paragraph that starts with a non-digit character,
        # contains Chinese text, and ends with "：" (looks like a section label).
        section_heading_pattern = re.compile(r"^[^\d\s][^：]*：")

        try:
            zin = zipfile.ZipFile(BytesIO(content), "r")
            sheet_xml = zin.read("word/document.xml")
            zin.close()
        except Exception:
            return content

        tree = etree.fromstring(sheet_xml)
        body = tree.find(f"{{{NS}}}body")
        if body is None:
            return content

        children = list(body)
        swapped = False

        i = 0
        while i < len(children) - 1:
            curr = children[i]
            nxt = children[i + 1]
            curr_tag = curr.tag.split("}")[1] if "}" in curr.tag else curr.tag
            nxt_tag = nxt.tag.split("}")[1] if "}" in nxt.tag else nxt.tag

            if curr_tag == "tbl" and nxt_tag == "p":
                # Extract text of next paragraph
                texts = [t.text or "" for t in nxt.findall(f".//{{{NS}}}t")]
                para_text = "".join(texts).strip()
                if section_heading_pattern.match(para_text):
                    # Swap: move heading before table
                    body.remove(nxt)
                    body.insert(i, nxt)
                    children = list(body)
                    swapped = True
                    # After swapping, re-evaluate at same position since a new
                    # element is now before the table; skip ahead to avoid infinite loop
                    i += 1
                    continue

            i += 1

        if not swapped:
            return content

        # Repack the docx with the reordered XML
        out = BytesIO()
        with zipfile.ZipFile(BytesIO(content), "r") as zin:
            with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.namelist():
                    if item == "word/document.xml":
                        zout.writestr(
                            item,
                            etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True),
                        )
                    else:
                        zout.writestr(item, zin.read(item))
        out.seek(0)
        return out.read()

    @staticmethod
    def _sanitize_doc_key(key: str) -> str:
        """Remove chars that OnlyOffice server-side routing mishandles."""
        import re
        return re.sub(r"[\(\)\s/\\]+", "_", key).strip("_")

    def _build_msds_docx_from_doc(self, doc_path: str, product_name: str) -> bytes:
        """
        将 .doc 文件转换为包含标准 MSDS 表格结构的 .docx。
        从源 .doc 提取文本填充到预设表格中。
        """
        text = self.msds_service.extract_text(doc_path)
        composition = self.msds_service.extract_composition_table(text)
        props = self.msds_service.extract_physical_props(text)

        doc = Document()

        # 标题
        title = doc.add_heading("物质安全数据表 (MSDS)", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 产品名
        doc.add_paragraph(f"产品名称：{product_name}").runs[0].bold = True
        doc.add_paragraph("")

        # 表1：理化特性
        doc.add_heading("1. 理化特性", level=2)
        table1 = doc.add_table(rows=4, cols=2)
        table1.style = "Table Grid"
        fields = [
            ("外观与性状", props.get("physical_form", "")),
            ("离子型", props.get("ion_type", "")),
            ("pH值", props.get("ph", "")),
            ("沸点/熔点", ""),
        ]
        for i, (label, value) in enumerate(fields):
            table1.rows[i].cells[0].text = label
            table1.rows[i].cells[1].text = value

        doc.add_paragraph("")

        # 表2：成分组成
        doc.add_heading("2. 成分组成", level=2)
        table2 = doc.add_table(rows=1 + len(composition), cols=3)
        table2.style = "Table Grid"
        # 表头
        hdr = table2.rows[0].cells
        hdr[0].text = "序号"
        hdr[1].text = "组分名称"
        hdr[2].text = "CAS No. / 含量"
        for cell in hdr:
            cell.paragraphs[0].runs[0].bold = True
        # 数据行
        for i, item in enumerate(composition):
            row = table2.rows[i + 1].cells
            row[0].text = str(i + 1)
            row[1].text = item.get("component", "")
            row[2].text = item.get("percentage", "")

        doc.add_paragraph("")

        # 附加信息（来自原文）
        doc.add_heading("3. 其他信息", level=2)
        for line in text.split("\n")[:50]:
            line = line.strip()
            if line and len(line) > 3:
                doc.add_paragraph(line)

        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()

    def generate_customs(
        self,
        order_id: int | None = None,       # OrderPiRecord 主键 ID（非业务订单号）
        ledger_record_id: int | None = None,
        order_no: str | None = None,
    ) -> Tuple[bytes, str, str]:
        """
        生成报关资料工作簿（5个 sheet 的 xlsx）。

        填充策略：
        - 报关单 sheet：完整填充（表头 + 产品明细 + 申报要素）
        - 其余 4 sheet：仅填充独立字段（发票编号/日期/付款条件等），
          产品数据由模板内 OFFSET 公式自动从报关单引用。

        数据来源优先级：
        1. ledger_record_id（台账记录）
        2. order_id（回退到 OrderPiRecord 查询）
        3. order_no（直接用于台账查询）
        """
        import openpyxl
        from app.services.ledger_service import LedgerService
        from app.services.customs_declaration_service import CustomsDeclarationService

        template_path = TEMPLATES["customs"]
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Customs template not found: {template_path}")

        # ── 读取台账数据 ────────────────────────────────────────────────
        if ledger_record_id:
            ledger_svc = LedgerService()
            record = ledger_svc.get_ledger_record(ledger_record_id)
        elif order_no:
            ledger_svc = LedgerService()
            resp = ledger_svc.list_ledger(search=order_no, page_size=1)
            record = resp.records[0] if resp.records else None
        elif order_id:
            db = SessionLocal()
            try:
                rows = db.query(OrderPiRecord).filter_by(id=order_id).all()
                if not rows:
                    record = None
                else:
                    from app.schemas.ledger import LedgerRecordResponse, LedgerItemSchema
                    first = rows[0]
                    record = LedgerRecordResponse(
                        id=first.id,
                        order_no=first.order_no,
                        customer_code=first.customer_code,
                        sales_person=first.sales_person,
                        consignee_name=first.consignee_name,
                        consignee_address=first.consignee_address,
                        destination=first.destination,
                        loading_port=first.loading_port,
                        price_term=first.price_term,
                        payment_terms=first.payment_terms,
                        pi_date=first.pi_date,
                        currency=first.currency,
                        items=[
                            LedgerItemSchema(
                                internal_code=r.internal_code,
                                product_cn=r.product_cn,
                                spec_kg=r.spec_kg,
                                quantity_kg=r.quantity_kg,
                                unit_price=r.unit_price,
                                total_amount=r.total_amount,
                                hs_code=r.hs_code,
                                customs_name=r.customs_name,
                                customs_ingredients=r.components,
                                product_appearance=r.product_appearance,
                                gross_weight_kg=r.gross_weight_kg,
                                net_weight_kg=r.net_weight_kg,
                                drum_count=r.drum_count,
                                pallet_count=r.pallet_count,
                            )
                            for r in rows
                        ],
                    )
            finally:
                db.close()
        else:
            record = None

        # ── 加载工作簿 ────────────────────────────────────────────────
        wb = openpyxl.load_workbook(template_path)
        ws_customs = wb["报关单"]
        ws_invoice = wb["发票"]
        ws_packing = wb["箱单"]
        ws_contract = wb["合同"]
        ws_proxy = wb["委托书"]

        decl_svc = CustomsDeclarationService.get_instance()

        # ── 通用辅助函数 ──────────────────────────────────────────────
        def cn_port(dest: str) -> str:
            """将英文港口/国家名翻译为中文海关常用名"""
            mapping = {
                "thailand": "泰国", "singapore": "新加坡",
                "shanghai": "上海", "guangzhou": "广州", "shenzhen": "深圳",
                "hong kong": "香港", "hkg": "香港",
                "malaysia": "马来西亚", "indonesia": "印度尼西亚",
                "jakarta": "雅加达", "indian": "印度", "india": "印度",
                "ahmedabad": "艾哈迈达巴德",
                "vietnam": "越南", "ho chi minh": "胡志明市", "hcmc": "胡志明市",
                "bangladesh": "孟加拉国", "philippines": "菲律宾", "manila": "马尼拉",
                "korea": "韩国", "busan": "釜山", "japan": "日本", "tokyo": "东京",
                "taiwan": "台湾", "taipei": "台北",
                "usa": "美国", "united states": "美国",
                "uk": "英国", "germany": "德国", "france": "法国",
                "netherlands": "荷兰", "italy": "意大利", "spain": "西班牙",
                "australia": "澳大利亚",
                "lat krabang": "拉格拉邦", "laem chabang": "拉格拉邦",
                "bangkok": "曼谷", "tangkok": "曼谷",
                "singapore port": "新加坡", "port of singapore": "新加坡",
                "shanghai port": "上海", "guangzhou port": "广州",
                "yantian": "盐田", "shekou": "蛇口", "chiwan": "赤湾",
                "ningbo": "宁波", "qingdao": "青岛", "tianjin": "天津",
                "dalian": "大连", "xiamen": "厦门",
            }
            if not dest:
                return ""
            lower = dest.lower().strip()
            if lower in mapping:
                return mapping[lower]
            for key, val in mapping.items():
                if key in lower or lower in key:
                    return val
            return dest

        # ── 无数据：返回未填充模板 ──────────────────────────────────
        if not record:
            buf = BytesIO()
            wb.save(buf)
            content = buf.getvalue()
            doc_key = f"customs_{uuid.uuid4().hex}"
            return content, doc_key, base64.b64encode(content).decode()

        items = record.items or []

        # ── 基本信息（从台账提取）──────────────────────────────────
        pi_no = record.order_no or ""
        consignee = record.consignee_name or ""
        consignee_addr = record.consignee_address or ""
        dest_raw = record.destination or ""
        dest_cn = cn_port(dest_raw)
        dest_port = dest_raw
        price_term = record.price_term or ""
        payment_terms = record.payment_terms or ""
        pi_date = record.pi_date or ""
        currency = getattr(record, "currency", None) or "CNY"  # 默认人民币
        total_pallets = sum(it.pallet_count or 0 for it in items)
        total_drums = sum(it.drum_count or 0 for it in items)
        total_gw = sum(it.gross_weight_kg or 0 for it in items)
        # 净重 = 产品本身重量（quantity_kg），不含包装
        total_nw = sum(it.quantity_kg or it.net_weight_kg or 0 for it in items)
        # 件数：有托盘用托盘数，无托盘用桶数/包数
        piece_count = total_pallets if total_pallets else total_drums

        # ══════════════════════════════════════════════════════════════
        # Sheet 1: 报关单 — 占位符替换
        # ══════════════════════════════════════════════════════════════
        ws = ws_customs
        # 固定值：A4 公司名（模板中已预填，保持不变）
        # 表头动态字段
        replace_placeholder(ws, "{{CONSIGNEE}}", consignee)
        replace_placeholder(ws, "{{CONTRACT_NO}}", pi_no)
        replace_placeholder(ws, "{{DEST_COUNTRY_CN}}", dest_cn)
        replace_placeholder(ws, "{{DEST_COUNTRY_CN_2}}", dest_cn)
        replace_placeholder(ws, "{{DEST_PORT}}", dest_port)
        replace_placeholder(ws, "{{TOTAL_PIECES}}", piece_count)
        replace_placeholder(ws, "{{TOTAL_GROSS_WEIGHT}}", round(total_gw, 1) if total_gw else None)
        replace_placeholder(ws, "{{TOTAL_NET_WEIGHT}}", round(total_nw, 1) if total_nw else None)
        replace_placeholder(ws, "{{PRICE_TERM}}", price_term)
        # V4: 成交方式（发票公式 =报关单!V4 引用）
        replace_placeholder(ws, "{{PRICE_TERM_V4}}", price_term)
        # 币制：I22 单元格
        replace_placeholder(ws, "{{CURRENCY}}", currency)

        # 产品明细：每品占 3 行（数据行 / 申报要素行 / 公式行）
        # 起始行 = 20 + idx * 3，模板最多支持到行 37（约 5-6 个产品）
        max_items = (ws.max_row - 20) // 3
        if len(items) > max_items:
            import logging
            logging.warning(
                "报关单模板最多容纳 %d 个产品，当前 %d 个，多余产品将被截断",
                max_items, len(items),
            )
        for idx, item in enumerate(items):
            row = 20 + idx * 3
            if row + 2 > ws.max_row:
                break

        # 业务说明：以下硬编码值适用于宏昊化工出口业务场景
        # 千克（KG）— 化工品标准计量单位
        # 中国 — 出口报关原产国
        # 肇庆 — 公司所在地，境内货源地
        # 照章征税 — 一般贸易默认征免方式

            ws.cell(row, 1, idx + 1)                           # A: 项号
            if item.hs_code:
                ws.cell(row, 2, item.hs_code)                  # B: 商品编号
            if item.customs_name:
                ws.cell(row, 4, item.customs_name)             # D: 商品名称
            if item.quantity_kg:
                ws.cell(row, 7, item.quantity_kg)              # G: 数量
            ws.cell(row, 8, "千克")                             # H: 单位
            if item.unit_price:
                ws.cell(row, 9, item.unit_price)               # I: 单价
            ws.cell(row, 11, "中国")                            # K: 原产国
            if dest_cn:
                ws.cell(row, 13, dest_cn)                      # M: 最终目的国
            ws.cell(row, 16, "肇庆")                            # P: 境内货源地
            ws.cell(row, 19, "照章征税")                        # S: 征免

            # 申报要素（行 row+1, D 列）— 从 JSON 复制整块，替换成分
            # 获取申报要素字符串
            elements_str = decl_svc.get_elements_str(item.hs_code or "", item.customs_name or "")

            if elements_str:
                # 成分：优先台账 customs_ingredients，换行替换为双空格
                ingredient = (item.customs_ingredients or "").replace("\n", "  ")
                if ingredient:
                    # 替换成分字段
                    decl_str = CustomsDeclarationService.replace_ingredient(elements_str, ingredient)
                else:
                    # 没有台账成分，使用 JSON 中的原始成分
                    decl_str = elements_str
            else:
                # JSON 中没有找到，手动构建
                parts = []
                if item.customs_name:
                    parts.append(f"用途：{item.customs_name}")
                if item.customs_ingredients:
                    parts.append(f"成分：{item.customs_ingredients.replace(chr(10), '  ')}")
                decl_str = "|".join(parts)

            if decl_str:
                ws.cell(row + 1, 4, decl_str)                  # D21: 申报要素

            # 总价（行 row+1, I 列）— 模板公式 =I20*G20 会自动计算，
            # 但 openpyxl 保存后公式缓存值可能不更新，所以直接写值
            if item.total_amount:
                ws.cell(row + 1, 9, item.total_amount)         # I21: 总价

        # ── T/U/V 辅助列：供箱单公式引用 ──
        # 箱单公式: OFFSET(报关单!$V$1, ROW(报关单!V{N})*3+16, 0)
        # 产品2: V23/U23/T23, 产品3: V26/U26/T26, ...
        for idx, item in enumerate(items):
            if idx == 0:
                continue  # 第1个产品由箱单直接引用 F12/G20
            helper_row = 23 + (idx - 1) * 3  # 23, 26, 29, 32, 35
            if helper_row > ws.max_row:
                break
            nw = item.quantity_kg or item.net_weight_kg
            gw = item.gross_weight_kg
            pc = item.pallet_count or item.drum_count
            if nw:
                ws.cell(helper_row, 20, round(nw, 1))   # T: 净重
            if gw:
                ws.cell(helper_row, 21, round(gw, 1))   # U: 毛重
            if pc:
                ws.cell(helper_row, 22, pc)               # V: 件数

        # ══════════════════════════════════════════════════════════════
        # Sheet 2: 发票 — 填充所有产品行
        # ══════════════════════════════════════════════════════════════
        ws = ws_invoice
        # 发票号：去掉合同号的字母前缀（HT/HH/HTPK等），换成 IN
        if pi_no:
            inv_no = _to_invoice_no(pi_no)
            ws["G2"] = inv_no
        if pi_date:
            ws["G3"] = _parse_date(pi_date)
        # G6 成交方式由公式 =报关单!V4 自动填充
        # 币制显示：USD → USD，CNY/RMB → 人民币
        currency_label = "人民币" if currency in ("CNY", "RMB", None) else currency
        # 填充所有产品行（行8开始，每品1行，完整填充所有列）
        total_amt_inv = 0
        for idx, item in enumerate(items):
            r = 8 + idx
            if r > 23:
                break
            ws.cell(r, 1).value = "N/M"                       # A: 唛头
            ws.cell(r, 3).value = item.customs_name or ""     # C: 货物名称
            ws.cell(r, 4).value = item.quantity_kg or 0       # D: 数量
            ws.cell(r, 5).value = "千克"                      # E: 单位
            ws.cell(r, 6).value = round(item.unit_price or 0, 2)  # F: 单价
            ws.cell(r, 7).value = currency_label              # G: 币制
            amt = round(item.total_amount or 0, 2)
            ws.cell(r, 8).value = amt                         # H: 总金额
            total_amt_inv += amt
        # 汇总行：直接写值（不依赖公式，避免缓存值问题）
        total_words = _amount_to_chinese_upper(total_amt_inv)
        ws["C24"].value = total_words                         # C: 中文大写金额
        ws["G24"].value = currency_label                      # G: 币制
        ws["H24"].value = round(total_amt_inv, 2)             # H: 数字金额
        # 备用：替换模板占位符
        replace_placeholder(ws, "{{TOTAL_AMOUNT_WORDS}}", total_words)
        replace_placeholder(ws, "{{TOTAL_AMOUNT}}", round(total_amt_inv, 2))
        replace_placeholder(ws, "{{CURRENCY}}", currency_label)

        # ── T/U/V 辅助列：供箱单公式引用 ──
        # 箱单公式: OFFSET(报关单!$V$1, ROW(报关单!V{N})*3+16, 0)
        # V4→V29, V5→V32, V6→V35, V7→V38, V8→V41
        ws = ws_customs  # 切回报关单
        for idx, item in enumerate(items):
            if idx == 0:
                continue  # 第1个产品由箱单直接引用 F12/G20
            # 箱单公式 ROW(V{4+idx})*3+16 = (4+idx)*3+16, 再+1偏移
            helper_row = 26 + idx * 3  # 29, 32, 35, 38, 41
            if helper_row > ws.max_row:
                break
            nw = item.quantity_kg or item.net_weight_kg or 0
            gw = item.gross_weight_kg or 0
            pc = item.pallet_count or item.drum_count or 0
            ws.cell(helper_row, 20).value = round(nw, 1)   # T: 净重
            ws.cell(helper_row, 21).value = round(gw, 1)   # U: 毛重
            ws.cell(helper_row, 22).value = pc               # V: 件数

        # ══════════════════════════════════════════════════════════════
        # Sheet 3: 箱单 — 填充所有产品行
        # ══════════════════════════════════════════════════════════════
        ws = ws_packing
        # 先清除所有数据行（包括模板残留的公式和空行）
        for r in range(10, 26):
            for c in range(1, 9):
                ws.cell(r, c).value = None
        # 行7：船名/目的地 — C7 显示 "广州 至 {destination}"
        if dest_raw:
            ws.cell(7, 3).value = f"广州  至  {dest_raw}"    # C7: 船名+目的地
        # 填充产品数据
        for idx, item in enumerate(items):
            if idx == 0:
                # 产品1: 行10
                r = 10
            else:
                # 产品2+: 行11, 12, 13, ...
                r = 10 + idx
            if r > 25:
                break
            pc = item.pallet_count or item.drum_count or 0
            qty = item.quantity_kg or 0
            gw = item.gross_weight_kg or 0
            nw = item.quantity_kg or item.net_weight_kg or 0
            unit = "托" if item.pallet_count else "桶"
            qty_unit = "千克"
            ws.cell(r, 1).value = "N/M" if idx == 0 else None   # A: 箱号（仅第1行）
            ws.cell(r, 2).value = item.customs_name or ""       # B: 货物名称
            ws.cell(r, 3).value = pc                            # C: 件数
            ws.cell(r, 4).value = unit                          # D: 件数单位
            ws.cell(r, 5).value = qty                           # E: 数量
            ws.cell(r, 6).value = qty_unit                      # F: 数量单位
            ws.cell(r, 7).value = round(gw, 1) if gw else None  # G: 毛重
            ws.cell(r, 8).value = round(nw, 1) if nw else None  # H: 净重

        # ══════════════════════════════════════════════════════════════
        # Sheet 4: 合同 — 填充占位符 + 无公式的独立单元格
        # （产品数据通过公式引用发票，无需手动填充）
        # ══════════════════════════════════════════════════════════════
        ws = ws_contract
        # 卖方名称（C2 已有公式 =报关单!A4，无需额外填充）
        # 买方名称（C8 已有公式 =报关单!A6，无需额外填充）
        if consignee_addr:
            replace_placeholder(ws, "{{CONSIGNEE_ADDR}}", consignee_addr)
            replace_placeholder(ws, "{{BUYER_ADDR}}", consignee_addr)
        # 固定值：G10 肇庆（模板中已预填，保持不变）
        # 成交方式：从PI数据填充（CIF/EXW/FOB/C&F等）
        replace_placeholder(ws, "{{PRICE_TERM}}", price_term)
        replace_placeholder(ws, "{{DEST_PORT}}", dest_port)
        replace_placeholder(ws, "{{PRICE_TERM_V4}}", price_term)
        # 日期：使用 PI 日期
        if pi_date:
            ws["G7"] = _parse_date(pi_date)
        # (9)装运口岸和目的地：广州---{destination}
        if dest_raw:
            ws["B39"] = f"(9)装运口岸和目的地           广州---{dest_raw}"
        # 合同总值大写
        total_amt_contract = sum(round(it.total_amount or 0, 2) for it in items)
        if total_amt_contract > 0:
            ws["D33"] = _amount_to_chinese_upper(total_amt_contract)

        # ══════════════════════════════════════════════════════════════
        # Sheet 5: 委托书 — 填充货物名称和HS编码（去重拼接）
        # ══════════════════════════════════════════════════════════════
        ws = ws_proxy
        # 货物名称：去重拼接
        all_names = list(dict.fromkeys(i.customs_name for i in items if i.customs_name))
        replace_placeholder(ws, "{{PROXY_GOODS_NAME}}", " / ".join(all_names) if all_names else None)
        # HS编码：去重拼接
        all_hs = list(dict.fromkeys(i.hs_code for i in items if i.hs_code))
        replace_placeholder(ws, "{{PROXY_HS_CODE}}", " / ".join(all_hs) if all_hs else None)

        # ── 保存 ────────────────────────────────────────────────
        buf = BytesIO()
        wb.save(buf)
        content = buf.getvalue()
        doc_key = f"customs_{uuid.uuid4().hex}"
        return content, doc_key, base64.b64encode(content).decode()

    def generate_template_instance(self, template_type: str) -> Tuple[bytes, str, str]:
        """加载空白模板（不填充 marker），返回 (content, doc_key, encoded_content)。"""
        type_map = {
            "booking": ("booking", "xlsx"),
            "msds":    ("msds",    "docx"),
        }
        if template_type not in type_map:
            raise ValueError(f"Unknown template type: {template_type}")
        key_prefix, file_ext = type_map[template_type]
        template_path = TEMPLATES[template_type]
        # Find effective template: .docx directly, or .xls for booking
        effective_template: Optional[str] = None
        if os.path.exists(template_path):
            if template_path.endswith(".docx"):
                effective_template = template_path
            elif template_path.endswith(".xls"):
                effective_template = template_path  # .xls will be converted to .xlsx below
        # For MSDS, also scan MSDS_DIR as fallback
        if not effective_template and template_type == "msds":
            for fname in os.listdir(MSDS_DIR):
                candidate = os.path.join(MSDS_DIR, fname)
                if fname.endswith(".docx") and os.path.exists(candidate):
                    effective_template = candidate
                    break
                elif fname.endswith(".doc") and os.path.exists(candidate):
                    try:
                        converted = convert_doc_to_docx(candidate)
                        temp_path = os.path.join(MSDS_DIR, f"_blank_converted_{fname}.docx")
                        with open(temp_path, "wb") as f:
                            f.write(converted)
                        effective_template = temp_path
                        break
                    except Exception:
                        continue
        if not effective_template:
            raise FileNotFoundError(f"Template not found: {template_path}")

        if file_ext == "xlsx":
            # .xls 模板 → .xlsx
            content = convert_xls_to_xlsx(effective_template)
        else:
            doc = Document(effective_template)
            buf = BytesIO()
            doc.save(buf)
            content = buf.getvalue()

        timestamp = int(time.time())
        doc_key = f"{key_prefix}_template_{timestamp}"
        return content, doc_key, base64.b64encode(content).decode()

    def load_msds_file(self, msds_index_record) -> Tuple[bytes, str, str]:
        """
        根据 MSDSIndex 记录加载原始文件，检测文件真实格式后转换。
        msds_index_record: MSDSIndex ORM 对象
        返回 (content_bytes, doc_key, base64_content)
        """
        file_path = msds_index_record.file_path

        # 用魔数检测真实格式，不用扩展名
        with open(file_path, "rb") as f:
            magic = f.read(8)

        # OLE2 Compound Document (.doc/.xls 等旧 Office 格式)
        if magic[:4] == b"\xd0\xcf\x11\xe0":
            content = convert_doc_to_docx(file_path)
        # PDF
        elif magic[:4] == b"%PDF":
            with open(file_path, "rb") as f:
                content = f.read()
        # ZIP-based docx/xlsx (默认)
        else:
            with open(file_path, "rb") as f:
                content = f.read()
            # Fix table/heading order in .docx before returning
            if file_path.endswith(".docx"):
                content = self._swap_table_heading_order(content)

        # 构建 doc_key（使用原始文件名，去掉空格括号等）
        import re
        safe_name = re.sub(r"[\(\)\s/\\]+", "_", os.path.splitext(os.path.basename(file_path))[0]).strip("_")
        doc_key = f"msds_{msds_index_record.id}_{safe_name}_{int(time.time())}"
        return content, doc_key, base64.b64encode(content).decode()
